#!/usr/bin/env python3
"""Reader-facing presentation contract for side stories.

Colour is redundant with a stable symbol + label. Multi-paragraph side stories receive
explicit top/left/bottom boundaries so the reader can see where the excursion ends.
New fenced/from-scratch renderers may use the same palette inside a single container.

Run27 contract: a `return_to` that is an artefact ID is resolved only from an explicit
canonical Markdown marker such as `[claim:C-...]`. We never guess the target from
semantic similarity or silently treat an ID as visible prose.
"""
from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from side_story_contract import KINDS, RENDER_LABELS, SIDE_STORY_CLASS, load_side_stories


KIND_ORDER = [
    "false_lead","detour","dezoom","also","method","portrait","object_focus","comparator","callback","analytical_focus",
]

SIDE_STORY_PRESENTATION = {
    "false_lead": {"symbol": "①", "fill": "FFF6D8", "border": "D6B656", "usage": "Question intuitive à tester puis corriger."},
    "detour": {"symbol": "②", "fill": "E3EEF7", "border": "7FA6C4", "usage": "Détour bref qui éclaire le fil principal."},
    "dezoom": {"symbol": "③", "fill": "E7EBF0", "border": "8797A8", "usage": "Changement d’échelle géographique ou historique."},
    "also": {"symbol": "④", "fill": "E7F3E8", "border": "83AA88", "usage": "Dimension complémentaire utile au mécanisme principal."},
    "method": {"symbol": "⑤", "fill": "F0EDE7", "border": "A69A88", "usage": "Méthode historique lue par le lecteur : dater, attribuer, comparer une source."},
    "portrait": {"symbol": "⑥", "fill": "F7E7E1", "border": "C88F7A", "usage": "Personnage replacé dans son rôle historique."},
    "object_focus": {"symbol": "⑦", "fill": "F5EFDD", "border": "B8A56C", "usage": "Objet, inscription, monument ou indice de terrain."},
    "comparator": {"symbol": "⑧", "fill": "EEE7F5", "border": "9D87B1", "usage": "Comparaison contrôlée avec limites explicites."},
    "callback": {"symbol": "⑨", "fill": "FBE6E3", "border": "C9877E", "usage": "Fil rouge : reprise narrative d’un fait déjà noué."},
    "analytical_focus": {"symbol": "⑩", "fill": "DCE6EE", "border": "6F8FA8", "usage": "Question analytique, contrastes, mécanismes et payoff."},
}

INK = RGBColor(32,55,72)
MUTED = RGBColor(92,99,108)
LEGEND_HEADING = "Légende des encadrés"
CANONICAL_TARGET_MARKER = re.compile(r"\[(?:claim|bridge|arc):([^\]]+)\]", re.I)
ID_LIKE = re.compile(r"^[A-Z][A-Z0-9]*-[A-Z0-9-]+$", re.I)


def _norm(value:str)->str:
    value=re.sub(r"^[①②③④⑤⑥⑦⑧⑨⑩]\s*","",value or "")
    value=unicodedata.normalize("NFKD",value)
    value="".join(ch for ch in value if not unicodedata.combining(ch))
    return re.sub(r"\s+"," ",value.casefold()).strip()


def _visible_markdown_text(value:str)->str:
    value=CANONICAL_TARGET_MARKER.sub("", value or "")
    value=re.sub(r"<!--.*?-->"," ",value,flags=re.S)
    value=re.sub(r"^#{1,6}\s+","",value.strip())
    value=re.sub(r"[*_`]","",value)
    value=re.sub(r"\[([^\]]+)\]\([^\)]+\)",r"\1",value)
    return re.sub(r"\s+"," ",value).strip()


def build_return_marker_map(canonical_markdown:str|None)->dict[str,str]:
    """Map artefact IDs to the reader-visible paragraph carrying their marker.

    The map is deterministic and marker-driven. Multiple markers may legitimately point
    to the same paragraph. Empty/marker-only blocks are ignored.
    """
    if not canonical_markdown:
        return {}
    mapping:dict[str,str]={}
    blocks=re.split(r"\n\s*\n",canonical_markdown)
    for block in blocks:
        ids=CANONICAL_TARGET_MARKER.findall(block)
        if not ids:
            continue
        visible=_visible_markdown_text(block)
        if not visible:
            continue
        for target_id in ids:
            mapping.setdefault(target_id,visible)
    return mapping


def style_name(kind:str)->str:return f"Side Story — {kind}"


def detect_kind(text:str)->str|None:
    normalized=_norm(text)
    for kind in KIND_ORDER:
        if normalized.startswith(_norm(RENDER_LABELS[kind])):return kind
    return None


def _remove_children(parent,tag:str)->None:
    for node in list(parent.findall(qn(tag))):parent.remove(node)


def _edge(name:str,color:str,size="8",space="0"):
    node=OxmlElement(f"w:{name}");node.set(qn("w:val"),"single");node.set(qn("w:sz"),size);node.set(qn("w:color"),color);node.set(qn("w:space"),space);return node


def _decorate_ppr(p_pr,kind:str,position:str="single")->None:
    """position: single | start | middle | end."""
    presentation=SIDE_STORY_PRESENTATION[kind]
    _remove_children(p_pr,"w:shd");_remove_children(p_pr,"w:pBdr")
    shd=OxmlElement("w:shd");shd.set(qn("w:fill"),presentation["fill"]);p_pr.append(shd)
    borders=OxmlElement("w:pBdr")
    borders.append(_edge("left",presentation["border"],size="18",space="8"))
    borders.append(_edge("right",presentation["border"],size="6",space="2"))
    if position in {"single","start"}:borders.append(_edge("top",presentation["border"],size="6",space="2"))
    if position in {"single","end"}:borders.append(_edge("bottom",presentation["border"],size="6",space="2"))
    p_pr.append(borders)


def ensure_side_story_styles(doc:Document)->None:
    missing=set(KINDS)-set(SIDE_STORY_PRESENTATION);extra=set(SIDE_STORY_PRESENTATION)-set(KINDS)
    if missing or extra:raise RuntimeError(f"side-story presentation drift: missing={sorted(missing)}, extra={sorted(extra)}")
    for kind in KIND_ORDER:
        name=style_name(kind);style=doc.styles[name] if name in doc.styles else doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        style.font.name="Calibri";style.font.size=Pt(10.5);style.font.color.rgb=INK
        style.paragraph_format.left_indent=Inches(0.18);style.paragraph_format.right_indent=Inches(0.12);style.paragraph_format.space_before=Pt(5);style.paragraph_format.space_after=Pt(6);style.paragraph_format.line_spacing=1.18;style.paragraph_format.keep_together=True
        _decorate_ppr(style._element.get_or_add_pPr(),kind,"single")


def _decorate_paragraph(paragraph,kind:str,*,header:bool,position:str="single")->None:
    _decorate_ppr(paragraph._p.get_or_add_pPr(),kind,position)
    if header:
        paragraph.style=style_name(kind);symbol=SIDE_STORY_PRESENTATION[kind]["symbol"]
        if not paragraph.text.lstrip().startswith(symbol):
            target=next((run for run in paragraph.runs if run.text),None)
            if target is not None:target.text=f"{symbol} {target.text}";target.bold=True
            else:run=paragraph.add_run(f"{symbol} ");run.bold=True


def _anchor_probe(text:str)->str:
    words=_norm(text).split()
    return " ".join(words[:18])


def _find_return_index(paragraphs,start:int,return_to:str|None,marker_map:dict[str,str]|None=None)->int|None:
    if not return_to:return None
    marker_map=marker_map or {}
    target=return_to.split(":",1)[1] if return_to.startswith("anchor:") else return_to
    # Explicit artefact IDs are never compared literally to reader prose. They must
    # resolve through a canonical marker carried by the Markdown source.
    if ID_LIKE.match(target):
        visible=marker_map.get(target)
        if not visible:
            return None
        normalized_target=_anchor_probe(visible)
    else:
        normalized_target=_norm(target)
    if not normalized_target:return None
    for index in range(start+1,min(len(paragraphs),start+80)):
        paragraph_text=_norm(paragraphs[index].text)
        if normalized_target in paragraph_text or paragraph_text in normalized_target:
            return index
    return None


def _story_header_index(paragraphs,item:dict,used:set[int])->int|None:
    kind=item.get("kind");label=_norm(RENDER_LABELS.get(kind,""));aliases=[item.get("title","")]+list((item.get("content") or {}).get("legacy_titles") or []);aliases=[_norm(alias) for alias in aliases if _norm(alias)]
    for index,paragraph in enumerate(paragraphs):
        if index in used or detect_kind(paragraph.text)!=kind:continue
        text=_norm(paragraph.text)
        if label and any(alias in text for alias in aliases):return index
    return None


def apply_side_story_palette(doc:Document,project:Path|None=None,canonical_markdown:str|None=None)->dict:
    """Apply full pastel blocks when a deterministic return target exists.

    ID-based targets resolve only through `[claim:ID]` / `[bridge:ID]` / `[arc:ID]`
    markers in canonical Markdown. Unresolved/single-paragraph stories still receive a
    visibly closed header box rather than a guessed range.
    """
    ensure_side_story_styles(doc);paragraphs=doc.paragraphs;header_kinds={};marker_map=build_return_marker_map(canonical_markdown)
    for index,paragraph in enumerate(paragraphs):
        kind=detect_kind(paragraph.text)
        if kind:header_kinds[index]=kind

    resolved_ranges:dict[int,tuple[int,str]]={}
    if project is not None:
        used_headers=set()
        for _,item in load_side_stories(project):
            if item.get("class")!=SIDE_STORY_CLASS or item.get("status") not in {"validated","promoted"}:continue
            kind=item.get("kind")
            if kind not in SIDE_STORY_PRESENTATION:continue
            start=_story_header_index(paragraphs,item,used_headers)
            if start is None:continue
            used_headers.add(start);end=_find_return_index(paragraphs,start,(item.get("placement") or {}).get("return_to"),marker_map)
            if end is not None and end>start+1:resolved_ranges[start]=(end,kind)

    body_count=0;resolved_blocks=0
    for start,kind in header_kinds.items():
        if start not in resolved_ranges:
            _decorate_paragraph(paragraphs[start],kind,header=True,position="single");continue
        end,_=resolved_ranges[start];stop=end
        for idx in range(start+1,end):
            if idx in header_kinds:stop=idx;break
        last=stop-1
        if last<=start:
            _decorate_paragraph(paragraphs[start],kind,header=True,position="single");continue
        resolved_blocks+=1;_decorate_paragraph(paragraphs[start],kind,header=True,position="start")
        for idx in range(start+1,stop):
            _decorate_paragraph(paragraphs[idx],kind,header=False,position="end" if idx==last else "middle");body_count+=1

    return {"headers_styled":len(header_kinds),"body_paragraphs_styled":body_count,"resolved_blocks":resolved_blocks,"kinds_seen":sorted(set(header_kinds.values())),"return_markers":len(marker_map)}


def _shade_cell(cell,fill:str)->None:
    tc_pr=cell._tc.get_or_add_tcPr();shd=tc_pr.find(qn("w:shd"))
    if shd is None:shd=OxmlElement("w:shd");tc_pr.append(shd)
    shd.set(qn("w:fill"),fill)


def style_side_story_cell(cell,kind:str)->None:
    """Use for fenced/from-scratch side stories: the whole excursion is one cell."""
    p=SIDE_STORY_PRESENTATION[kind];_shade_cell(cell,p["fill"]);tc_pr=cell._tc.get_or_add_tcPr();borders=tc_pr.find(qn("w:tcBorders"))
    if borders is None:borders=OxmlElement("w:tcBorders");tc_pr.append(borders)
    for edge in ("top","bottom","left","right"):
        borders.append(_edge(edge,p["border"],size="6" if edge!="left" else "18"))
    mar=tc_pr.find(qn("w:tcMar"))
    if mar is None:mar=OxmlElement("w:tcMar");tc_pr.append(mar)
    for name,value in (("top",120),("bottom",120),("start",160),("end",160)):
        node=OxmlElement(f"w:{name}");node.set(qn("w:w"),str(value));node.set(qn("w:type"),"dxa");mar.append(node)


def add_side_story_legend(doc:Document)->int:
    if any(p.text.strip()==LEGEND_HEADING for p in doc.paragraphs):return 0
    doc.add_page_break();heading=doc.add_heading(LEGEND_HEADING,level=1);heading.paragraph_format.keep_with_next=True
    p=doc.add_paragraph("La couleur facilite le repérage, mais le symbole et le libellé portent le sens en noir et blanc. Les numéros ne sont pas un ordre de priorité. « Point de méthode » décrit une méthode historique utile au lecteur ; il ne désigne jamais le processus de production de ce document.");p.paragraph_format.space_after=Pt(8)
    table=doc.add_table(rows=1,cols=3);table.autofit=False;headers=("Repère","Type d’encadré","Usage")
    for cell,text in zip(table.rows[0].cells,headers):
        cell.text=text;_shade_cell(cell,"E7EBF0")
        for run in cell.paragraphs[0].runs:run.bold=True;run.font.color.rgb=INK
    for kind in KIND_ORDER:
        row=table.add_row().cells;p=SIDE_STORY_PRESENTATION[kind];row[0].text=p["symbol"];row[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER;row[1].text=RENDER_LABELS[kind];row[2].text=p["usage"];_shade_cell(row[0],p["fill"]);_shade_cell(row[1],p["fill"])
    return len(KIND_ORDER)


def markdown_legend()->str:
    lines=[f"## {LEGEND_HEADING}","","La couleur facilite le repérage, mais le symbole et le libellé portent le sens en noir et blanc. Les numéros ne sont pas un ordre de priorité.",""]
    for kind in KIND_ORDER:
        p=SIDE_STORY_PRESENTATION[kind];lines.append(f"- {p['symbol']} **{RENDER_LABELS[kind]}** — {p['usage']}")
    lines.extend(["","**Point de méthode** désigne ici une méthode historique utile au lecteur ; il ne décrit jamais le processus de production, les runs, les statuts ou les versions du document."])
    return "\n".join(lines)
