#!/usr/bin/env python3
"""Render the reviewed, materialized V4 Markdown reader into DOCX.

The V4 manuscript is already the evidence-preserving composition product.  This
renderer is deliberately presentation-only: it does not rewrite prose or load a
different claim/source state.  It recognizes the legacy chapter headings and
pipe tables that predate the newer from-scratch renderer.
"""
from __future__ import annotations

import argparse
import html
import json
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from frontstage_reader_contract import assert_no_known_backstage_leak, visible_docx_text
from from_scratch_review_contract import SIDE_BEGIN, SIDE_END
from render_reader_exports import (
    INK,
    MUTED,
    add_hyperlink,
    add_inline_markdown,
    configure_section,
    configure_styles,
    set_cell_border,
    set_table_geometry,
    set_run_font,
    shade_cell,
)
from side_story_contract import RENDER_LABELS
from side_story_presentation import SIDE_STORY_PRESENTATION, detect_kind, style_side_story_cell


REPO = Path(__file__).resolve().parents[1]
PROJECT = REPO / "examples" / "sri_lanka_pre_1948"
OUTPUT = PROJECT / "09_output"
DEFAULT_SOURCE = OUTPUT / "report_v4_full.md"
DEFAULT_OUTPUT = OUTPUT / "Sri_Lanka_Fresque_historico_geographique_vol_retour_v4.docx"
POST_PROJECT = REPO / "examples" / "sri_lanka_post_1948"
POST_OUTPUT = POST_PROJECT / "09_output"
RENDER_SPECS = {
    "pre": {
        "output_dir": OUTPUT,
        "source": DEFAULT_SOURCE,
        "output": DEFAULT_OUTPUT,
        "header": "Sri Lanka · des origines à 1948 · V4",
        "eyebrow": "FRESQUE HISTORICO-GÉOGRAPHIQUE · VOLUME 1",
        "title": "Sri Lanka — une île construite par ses interfaces",
        "subtitle": "Des premiers réseaux de l’océan Indien à l’État colonial de 1948",
        "run_label": "Composition RUN47–RUN54",
        "date": "3 SEPTEMBRE 2026",
        "subject": "V4 matérialisée après la passe éditoriale RUN54",
        "keywords": "Sri Lanka, histoire, géographie, océan Indien, reader V4",
    },
    "post": {
        "output_dir": POST_OUTPUT,
        "source": POST_OUTPUT / "report_v4_full.md",
        "output": POST_OUTPUT / "Sri_Lanka_1948_2026_etude_historico_geographique_v4.docx",
        "header": "Sri Lanka · 1948 à 2026 · V4",
        "eyebrow": "FRESQUE HISTORICO-GÉOGRAPHIQUE · VOLUME 2",
        "title": "Sri Lanka — reconstruire l’État après 1948",
        "subtitle": "Continuités institutionnelles, guerres, développement et recomposition politique",
        "run_label": "Composition RUN56",
        "date": "3 SEPTEMBRE 2026",
        "subject": "V4 post-1948 matérialisée et relue chapitre par chapitre",
        "keywords": "Sri Lanka, histoire, géographie, indépendance, guerre civile, reader V4",
    },
}
GOLD = RGBColor(150, 111, 32)
BLUE = RGBColor(46, 116, 181)
TABLE_TOTAL_DXA = 9120


def clean_markup(text: str) -> str:
    text = re.sub(r"</?u>", "", text.strip())
    return re.sub(r"\\([\[\]])", r"\1", text)


def add_rich_text(paragraph, text: str) -> None:
    """Render ordinary emphasis plus readable/clickable Markdown links."""
    value = clean_markup(text)
    pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
    pos = 0
    for match in pattern.finditer(value):
        if match.start() > pos:
            add_inline_markdown(paragraph, value[pos:match.start()])
        add_hyperlink(paragraph, match.group(1), match.group(2))
        pos = match.end()
    if pos < len(value):
        add_inline_markdown(paragraph, value[pos:])


class ParagraphHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.paragraphs: list[str] = []
        self._buffer: list[str] = []
        self._in_p = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag == "p":
            self._in_p = True
            self._buffer = []

    def handle_data(self, data: str) -> None:
        if self._in_p:
            self._buffer.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag == "p" and self._in_p:
            value = html.unescape("".join(self._buffer)).strip()
            if value:
                self.paragraphs.append(value)
            self._in_p = False
            self._buffer = []


def html_paragraphs(raw: str) -> list[str]:
    parser = ParagraphHTMLParser()
    parser.feed(raw)
    return parser.paragraphs


def illustration_index(output_dir: Path = OUTPUT) -> dict[str, dict]:
    records: dict[str, dict] = {}
    for path in sorted((output_dir / "illustrations").glob("*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        items = data if isinstance(data, list) else [data]
        for item in items:
            marker = (item.get("render") or {}).get("marker")
            if marker:
                records[marker] = item
    return records


def add_cover(doc: Document, spec: dict | None = None) -> None:
    spec = spec or RENDER_SPECS["pre"]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(104)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(spec["eyebrow"]), size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(spec["title"]), size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    set_run_font(
        p.add_run(spec["subtitle"]),
        size=14,
        color=BLUE,
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4560, 4560])
    for cell in table.rows[0].cells:
        shade_cell(cell, "F4F6F9")
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": "D7DBE2"},
            bottom={"val": "single", "sz": "4", "color": "D7DBE2"},
        )
    left, right = table.rows[0].cells
    set_run_font(left.paragraphs[0].add_run("ÉDITION LECTEUR V4\n"), size=9, color=MUTED, bold=True)
    set_run_font(left.paragraphs[0].add_run(spec["run_label"]), size=9.5, color=INK)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.paragraphs[0].add_run(f"{spec['date']}\n"), size=9, color=MUTED, bold=True)
    set_run_font(right.paragraphs[0].add_run("Version matérialisée et vérifiée"), size=9.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("Une lecture chronologique guidée par les lieux, les mécanismes et leurs sources."),
        size=10,
        color=MUTED,
        italic=True,
    )
    doc.add_page_break()


def configure_v4_styles(doc: Document) -> None:
    configure_styles(doc)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20
    for name, size, before, after in (
        ("Heading 1", 17, 18, 9),
        ("Heading 2", 13.5, 11, 5),
        ("Heading 3", 11.5, 8, 4),
        ("Heading 4", 10.5, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def strip_heading_markup(text: str) -> str:
    return clean_markup(re.sub(r"^\*\*(.*?)\*\*$", r"\1", text.strip()))


def add_heading(doc: Document, text: str, level: int, *, page_break: bool = False) -> None:
    p = doc.add_heading(level=min(level, 4))
    if page_break:
        # A standalone break paragraph can itself be pushed to the next page
        # when the preceding page is full, yielding an unintended blank page.
        # Binding the break to the heading is stable in Word and LibreOffice.
        p.paragraph_format.page_break_before = True
    add_rich_text(p, strip_heading_markup(text))


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    if columns == 1:
        return [TABLE_TOTAL_DXA]
    weights = []
    for index in range(columns):
        longest = max(len(re.sub(r"[*_`\[\]()]", "", row[index])) for row in rows)
        weights.append(max(12, min(longest, 54)))
    total = sum(weights)
    widths = [round(TABLE_TOTAL_DXA * weight / total) for weight in weights]
    widths[-1] += TABLE_TOTAL_DXA - sum(widths)
    return widths


def is_separator(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in row)


def parse_table_line(line: str) -> list[str]:
    return [clean_markup(cell.strip()) for cell in line.strip().strip("|").split("|")]


def render_table(doc: Document, raw_lines: list[str]) -> None:
    rows = [parse_table_line(line) for line in raw_lines]
    if len(rows) >= 2 and is_separator(rows[1]):
        rows.pop(1)
        header = True
    else:
        header = False
    columns = len(rows[0])
    if any(len(row) != columns for row in rows):
        raise RuntimeError("inconsistent Markdown table geometry")
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_geometry(table, table_widths(rows))
    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        for cell, value in zip(row.cells, values):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            add_rich_text(p, value)
            set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": "D7DBE2"})
            if header and row_index == 0:
                shade_cell(cell, "E7EBF0")
                for run in p.runs:
                    run.bold = True
            elif columns == 1:
                shade_cell(cell, "F4F6F9")
                set_cell_border(
                    cell,
                    left={"val": "single", "sz": "18", "color": "2E74B5"},
                    top={"val": "single", "sz": "3", "color": "D7DBE2"},
                    bottom={"val": "single", "sz": "3", "color": "D7DBE2"},
                )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_side_story(
    doc: Document,
    kind: str,
    lines: list[str],
    illustrations: dict[str, dict] | None = None,
) -> int:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_TOTAL_DXA])
    cell = table.cell(0, 0)
    style_side_story_cell(cell, kind)
    cell.paragraphs[0].text = ""
    first = True
    illustration_captions = 0
    for raw in lines:
        text = raw.strip()
        if not text or text.startswith("<!--"):
            continue
        if re.fullmatch(r"\[ILLUSTRATION:[^\]]+\]", text):
            item = (illustrations or {}).get(text)
            if item and (item.get("render") or {}).get("required_in_reader"):
                # The reviewed Markdown already carries the human caption on the
                # following line.  Count the anchor but suppress the backstage ID;
                # adding the registry caption here would duplicate reader prose.
                illustration_captions += 1
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        heading = re.match(r"^(#{1,4})\s+(.+)$", text)
        body = heading.group(2) if heading else text
        if first:
            expected = RENDER_LABELS[kind]
            symbol = SIDE_STORY_PRESENTATION[kind]["symbol"]
            if expected.casefold() not in re.sub(r"[*_]", "", body).casefold():
                body = f"**{expected}** — {body}"
            body = re.sub(r"^\*\*", f"**{symbol} ", body, count=1)
            p.paragraph_format.keep_with_next = True
        add_rich_text(p, body)
        if heading:
            for run in p.runs:
                run.bold = True
        first = False
    if first:
        raise RuntimeError(f"empty side-story block: {kind}")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return illustration_captions


def add_html_callout(doc: Document, raw: str) -> None:
    paragraphs = html_paragraphs(raw)
    if not paragraphs:
        return
    kind = detect_kind(paragraphs[0])
    if kind:
        add_side_story(doc, kind, paragraphs)
        return
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_TOTAL_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": "2E74B5"})
    for index, value in enumerate(paragraphs):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_rich_text(p, value)
        if index < 2:
            for run in p.runs:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_illustration_caption(doc: Document, item: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_TOTAL_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, "F5EFDD")
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": "B8A56C"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("REPÈRE D’ILLUSTRATION — "), size=9.5, color=GOLD, bold=True)
    add_rich_text(p, (item.get("fragment") or {}).get("caption", "Photo de terrain associée."))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body_line(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped or stripped == "---" or stripped.startswith("<!--"):
        return
    chapter = re.fullmatch(r"\*\*Chapitre\s+\d+\s+—.*\*\*", stripped)
    if chapter:
        add_heading(doc, stripped, 1, page_break=True)
        return
    heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
    if heading:
        title = strip_heading_markup(heading.group(2))
        page_break = title.startswith(("Chapitre ", "Épilogue", "Annexe technique"))
        add_heading(doc, title, len(heading.group(1)), page_break=page_break)
        return
    if stripped.startswith("> "):
        p = doc.add_paragraph(style="Causal Callout")
        add_rich_text(p, stripped[2:])
        return
    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        add_rich_text(p, re.sub(r"^\d+\.\s+", "", stripped))
        return
    if stripped.startswith(("- ", "* ")):
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, stripped[2:])
        return
    p = doc.add_paragraph()
    add_rich_text(p, stripped)


def add_toc_line(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("# "):
        add_heading(doc, stripped[2:], 1)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if stripped.startswith("> "):
        p.paragraph_format.left_indent = Inches(0.22)
        stripped = stripped[2:]
    is_part = re.fullmatch(r"\*\*PARTIE\s+[IVX]+\s+—.*\*\*", stripped) is not None
    add_rich_text(p, stripped)
    for run in p.runs:
        run.font.size = Pt(9 if is_part else 8.5)
        if is_part:
            run.bold = True


def render_markdown(doc: Document, markdown: str, *, output_dir: Path = OUTPUT) -> dict[str, int]:
    lines = markdown.splitlines()
    index = 0
    side_stories = 0
    tables = 0
    html_callouts = 0
    illustration_captions = 0
    illustrations = illustration_index(output_dir)
    toc_active = False
    seen_parts: set[str] = set()
    while index < len(lines):
        stripped = lines[index].strip()
        part = re.fullmatch(r"\*\*(PARTIE\s+[IVX]+\s+—.*?)\*\*", stripped)
        if stripped == "# **Sommaire**":
            toc_active = True
            add_toc_line(doc, lines[index])
            index += 1
            continue
        if part:
            key = part.group(1)
            if key in seen_parts:
                # Part labels already appear in the table of contents.  Repeating
                # them alone before page-breaking chapter titles creates orphans.
                if key.startswith("PARTIE I "):
                    toc_active = False
                index += 1
                continue
            seen_parts.add(key)
        if toc_active:
            add_toc_line(doc, lines[index])
            index += 1
            continue
        begin = SIDE_BEGIN.match(stripped)
        if begin:
            story_id, kind = begin.groups()
            body = []
            index += 1
            while index < len(lines) and not SIDE_END.match(lines[index].strip()):
                body.append(lines[index])
                index += 1
            if index >= len(lines) or SIDE_END.match(lines[index].strip()).group(1) != story_id:
                raise RuntimeError(f"unclosed side-story block: {story_id}")
            illustration_captions += add_side_story(doc, kind, body, illustrations)
            side_stories += 1
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            render_table(doc, table_lines)
            tables += 1
            continue
        if stripped == "<table>":
            raw_html = []
            while index < len(lines):
                raw_html.append(lines[index])
                if lines[index].strip() == "</table>":
                    break
                index += 1
            if index >= len(lines) or lines[index].strip() != "</table>":
                raise RuntimeError("unclosed HTML callout table")
            add_html_callout(doc, "\n".join(raw_html))
            html_callouts += 1
            index += 1
            continue
        illustration = re.fullmatch(r"\[ILLUSTRATION:[^\]]+\]", stripped)
        if illustration:
            item = illustrations.get(stripped)
            if item and (item.get("render") or {}).get("required_in_reader"):
                add_illustration_caption(doc, item)
                illustration_captions += 1
            index += 1
            continue
        add_body_line(doc, lines[index])
        index += 1
    return {
        "side_story_blocks": side_stories,
        "tables": tables,
        "html_callouts": html_callouts,
        "illustration_captions": illustration_captions,
    }


def build(source: Path, output: Path, spec: dict | None = None) -> dict[str, int | str]:
    spec = spec or RENDER_SPECS["pre"]
    markdown = source.read_text(encoding="utf-8")
    doc = Document()
    configure_v4_styles(doc)
    configure_section(doc.sections[0], spec["header"])
    add_cover(doc, spec)
    metrics = render_markdown(doc, markdown, output_dir=Path(spec["output_dir"]))
    visible = visible_docx_text(doc)
    assert_no_known_backstage_leak(visible)
    for token in ("<table>", "</table>", "<colgroup>", "[ILLUSTRATION:", "<!-- [RUN"):
        if token in visible:
            raise RuntimeError(f"V4 frontstage leak: {token}")
    core = doc.core_properties
    core.title = spec["title"]
    core.subject = spec["subject"]
    core.author = "Projet tourisme-etude-historico-geographique"
    core.keywords = spec["keywords"]
    core.comments = "Rendered from report_v4_full.md; evidence and composition remain separate."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    try:
        rendered_path = str(output.relative_to(REPO))
    except ValueError:
        rendered_path = str(output)
    return {"docx": rendered_path, **metrics}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project", choices=("pre", "post", "all"), default="pre")
    parser.add_argument("--source", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    if args.project == "all" and (args.source or args.output):
        parser.error("--source/--output cannot be combined with --project all")
    selected = ("pre", "post") if args.project == "all" else (args.project,)
    for key in selected:
        spec = RENDER_SPECS[key]
        source = args.source or Path(spec["source"])
        output = args.output or Path(spec["output"])
        print(build(source, output, spec))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
