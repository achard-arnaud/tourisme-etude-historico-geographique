#!/usr/bin/env python3
"""Build lossless Sri Lanka V3 readers from the long V1 baselines plus Run 5 deltas."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document

from render_reader_exports import add_inline_markdown, collect_register, extract_source_ids
from side_story_contract import assert_rendered_side_stories, validate_or_raise
from build_story_scaffold import build_scaffold
from toc_contract import inject_word_toc


REPO = Path(__file__).resolve().parents[1]
WORD_RE = re.compile(r"\b[\wÀ-ÖØ-öø-ÿĀ-ž'’.-]+\b", re.UNICODE)


SPECS = {
    "pre": {
        "project": REPO / "examples/sri_lanka_pre_1948",
        "baseline": "archive/Sri_Lanka_Fresque_historico_geographique_vol_retour_v1.docx",
        "delta": "report.md",
        "markdown_baseline": "report_v1_full.md",
        "markdown_output": "report_v3_full.md",
        "output": "Sri_Lanka_Fresque_historico_geographique_vol_retour_v3.docx",
        "cover_replacements": {
            3: "ÉDITION V3 INTÉGRALE DE LECTURE — VOL RETOUR",
            4: "V1 longue conservée, fiches de conversation et ajouts promus intégrés sans compression.",
            5: "Usage personnel — Sri Lanka, 20 août 2026",
        },
        "method_anchor": "Comment lire cette fresque",
        "method_block": """## Politique éditoriale de la V3 intégrale

Cette édition prend la V1 de soixante et une pages comme baseline non destructible. Les fiches de conversation, détours de terrain, side stories tracées et ajouts stabilisés sont insérés dans leur séquence historique ; aucun budget de longueur ne peut justifier leur suppression. Les répétitions utiles, controverses, limites de source et bifurcations sont conservées pour un lectorat avancé.

Le petit `report.md` est traité comme un delta promu et non comme un nouveau manuscrit complet. Les side stories requises portent un marqueur de lineage vérifié mais ce marqueur technique reste invisible dans le DOCX.
""",
        "groups": [
            ("7. Jaffna, Mannar et le verrou du nord", [1]),
            ("3. La côte hollandaise : de la forteresse à la bureaucratie", [2, 3]),
            ("8. Kandy joue encore l’international", [4]),
            ("9. 1795–1796 — la VOC tombe à cause de l’Europe, pas de Kandy", [5]),
            ("12. Le grand paradoxe britannique", [6]),
        ],
        "source_anchor": "FIN DE L’ÉDITION DE LECTURE",
        "complement_heading_level": 3,
        "title": "Sri Lanka — enquête historico-géographique intégrale, des origines à 1948",
    },
    "post": {
        "project": REPO / "examples/sri_lanka_post_1948",
        "baseline": "archive/Sri_Lanka_1948_2026_etude_historico_geographique_v1.docx",
        "delta": "report.md",
        "markdown_baseline": "report_v1_full.md",
        "markdown_output": "report_v3_full.md",
        "output": "Sri_Lanka_1948_2026_etude_historico_geographique_v3.docx",
        "cover_replacements": {
            3: "V3 intégrale — arcs chronologiques, HIL, comparateurs et zooms géographiques",
            4: "État des données : 20 août 2026 · Usage personnel · Baseline V1 conservée",
        },
        "method_anchor": "Méthode de lecture",
        "method_block": """## Politique éditoriale de la V3 intégrale

Cette édition conserve la totalité de la V1 moderne et ajoute les développements promus sur langue, caste, éducation, guerre, diaspora, Tamil Nadu, Indonésie et conversion territoriale du capital humain. Pour ce lectorat avancé, aucun plafond de mots n'est appliqué ; la comparaison est développée avec ses confounders et ses limites.
""",
        "groups": [
            ("ARC A01 — 1948-1956", [1]),
            ("ARC A02 — 1956-1972", [2, 3, 9, 13]),
            ("ARC A03 — 1972-1983", [4, 5, 6]),
            ("ARC A04 — 1983-2002", [7, 8]),
            ("ARC A06 — 2009-2019", [10, 11, 12]),
            ("Synthèse — sept décennies", [14]),
        ],
        "source_anchor": "Sources d'ancrage — sélection",
        "complement_heading_level": 2,
        "title": "Sri Lanka 1948–2026 — étude historico-géographique intégrale",
    },
}


def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def docx_word_count(doc: Document) -> int:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return word_count("\n".join(chunks))


def numbered_sections(text: str) -> dict[int, str]:
    matches = list(re.finditer(r"(?m)^##\s+(\d+)\.\s+", text))
    sections: dict[int, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        sections[int(match.group(1))] = text[match.start():end].strip()
    return sections


def as_v3_complement(section: str, heading_level: int) -> str:
    """Demote a delta chapter so baseline chapter numbering remains authoritative."""
    return re.sub(
        r"(?m)^##\s+\d+\.\s+(.+)$",
        f"{'#' * heading_level} Complément V3 — \\1",
        section,
        count=1,
    )


def replace_preserving_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph(doc: Document, anchor: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(anchor):
            return paragraph
    raise ValueError(f"anchor not found in baseline DOCX: {anchor}")


def markdown_nodes(doc: Document, markdown: str):
    nodes = []
    for raw in markdown.splitlines():
        line = raw.strip()
        if (
            not line
            or line == "---"
            or line.startswith("> **Statut**")
            or (line.startswith("<!--") and line.endswith("-->"))
        ):
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            paragraph = doc.add_heading(level=level)
            add_inline_markdown(paragraph, heading.group(2))
        elif line.startswith("> "):
            style = "Quote" if "Quote" in [s.name for s in doc.styles] else None
            paragraph = doc.add_paragraph(style=style)
            add_inline_markdown(paragraph, line[2:])
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, line)
        elif line.startswith(("- ", "* ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:])
        else:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, line)
        nodes.append(paragraph._p)
    return nodes


def insert_after_docx_anchor(doc: Document, anchor: str, markdown: str) -> None:
    paragraph = find_paragraph(doc, anchor)
    nodes = markdown_nodes(doc, markdown)
    for node in reversed(nodes):
        paragraph._p.addnext(node)


def insert_before_docx_anchor(doc: Document, anchor: str, markdown: str, page_break=False) -> None:
    paragraph = find_paragraph(doc, anchor)
    nodes = markdown_nodes(doc, markdown)
    if page_break:
        nodes.insert(0, doc.add_page_break()._p)
    for node in nodes:
        paragraph._p.addprevious(node)


def normalize_heading(line: str) -> str:
    return re.sub(r"[*_`]", "", re.sub(r"^#+\s*", "", line)).strip()


def insert_after_markdown_anchor(text: str, anchor: str, block: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if normalize_heading(line).startswith(anchor):
            lines[index + 1:index + 1] = ["", block.strip(), ""]
            return "\n".join(lines)
    raise ValueError(f"anchor not found in baseline Markdown: {anchor}")


def insert_before_markdown_anchor(text: str, anchor: str, block: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if normalize_heading(line).startswith(anchor):
            lines[index:index] = ["", block.strip(), ""]
            return "\n".join(lines)
    raise ValueError(f"anchor not found in baseline Markdown: {anchor}")


def source_appendix(project: Path, delta: str) -> str:
    register = collect_register(project)
    source_ids = extract_source_ids(delta, register)
    lines = [
        "## Appareil de sources des compléments V3",
        "",
        "Ces notices donnent la source exacte des développements ajoutés à la V1. Le tier décrit la nature de la source, non une autorisation à généraliser au-delà de son champ.",
    ]
    for source_id in source_ids:
        source = register[source_id]
        lines.extend([
            "",
            f"### {source_id} — {source.get('title', '')}",
            " ".join(part for part in [source.get("tier", ""), source.get("anchor_role", ""), source.get("scope", "")] if part),
            f"Limite : {source.get('limitations', '')}",
            source.get("url", ""),
        ])
    return "\n".join(lines)


def build(key: str) -> dict[str, int | str]:
    spec = SPECS[key]
    project = spec["project"]
    output_dir = project / "09_output"
    baseline_path = output_dir / spec["baseline"]
    delta_path = output_dir / spec["delta"]
    markdown_baseline_path = output_dir / spec["markdown_baseline"]
    output_path = output_dir / spec["output"]
    markdown_output_path = output_dir / spec["markdown_output"]

    side_story_count = validate_or_raise(project, check_render=True)
    delta = delta_path.read_text(encoding="utf-8")
    sections = numbered_sections(delta)
    expected_sections = {number for _, numbers in spec["groups"] for number in numbers}
    missing = sorted(expected_sections - sections.keys())
    if missing:
        raise ValueError(f"missing numbered delta sections for {key}: {missing}")

    doc = Document(baseline_path)
    baseline_words = docx_word_count(doc)
    for index, replacement in spec["cover_replacements"].items():
        replace_preserving_first_run(doc.paragraphs[index], replacement)
    scaffold_path=output_dir/"story_scaffold.json"
    scaffold=json.loads(scaffold_path.read_text(encoding="utf-8")) if scaffold_path.exists() else build_scaffold(project)
    toc_arcs=inject_word_toc(doc,scaffold,max(spec["cover_replacements"]))

    insert_after_docx_anchor(doc, spec["method_anchor"], spec["method_block"])
    for anchor, numbers in spec["groups"]:
        block = "\n\n".join(as_v3_complement(sections[number], spec["complement_heading_level"]) for number in numbers)
        insert_after_docx_anchor(doc, anchor, block)
    appendix = source_appendix(project, delta)
    insert_before_docx_anchor(doc, spec["source_anchor"], appendix, page_break=True)

    doc.core_properties.title = spec["title"]
    doc.core_properties.subject = "V3 intégrale, lectorat avancé, sans plafond de longueur"
    doc.core_properties.comments = (
        "Built losslessly from the archived V1 reader plus the complete promoted delta; "
        "side-story lineage is validated before export; TOC hierarchy is scaffold-backed."
    )

    output_words = docx_word_count(doc)
    delta_words = word_count(delta)
    minimum_words = baseline_words + int(delta_words * 0.90)
    if output_words < minimum_words:
        raise RuntimeError(
            f"retention gate failed for {key}: {output_words} < {minimum_words} "
            f"(baseline={baseline_words}, delta={delta_words})"
        )
    doc.save(output_path)

    markdown = markdown_baseline_path.read_text(encoding="utf-8")
    markdown = insert_after_markdown_anchor(markdown, spec["method_anchor"], spec["method_block"])
    for anchor, numbers in spec["groups"]:
        block = "\n\n".join(as_v3_complement(sections[number], spec["complement_heading_level"]) for number in numbers)
        markdown = insert_after_markdown_anchor(markdown, anchor, block)
    markdown = insert_before_markdown_anchor(markdown, spec["source_anchor"], appendix)
    assert_rendered_side_stories(project, markdown)
    markdown_output_path.write_text(markdown.rstrip() + "\n", encoding="utf-8")

    return {
        "project": key,
        "baseline_docx_words": baseline_words,
        "delta_words": delta_words,
        "v3_docx_words": output_words,
        "retention_vs_baseline_percent": round(output_words / baseline_words * 100, 1),
        "side_stories": side_story_count,
        "toc_arcs":toc_arcs,
        "docx": str(output_path.relative_to(REPO)),
        "markdown": str(markdown_output_path.relative_to(REPO)),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project", choices=["pre", "post", "all"], default="all")
    args = parser.parse_args()
    keys = list(SPECS) if args.project == "all" else [args.project]
    metrics = [build(key) for key in keys]
    metrics_path = REPO / "docs/RUN7_V3_RETENTION_METRICS.json"
    metrics_path.write_text(json.dumps(metrics, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(metrics, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
