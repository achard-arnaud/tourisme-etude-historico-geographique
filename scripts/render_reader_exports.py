#!/usr/bin/env python3
"""Render deliberately abridged historical-travel Markdown as DOCX reader editions.

The lossless advanced Sri Lanka workflow is ``render_full_reader_v3.py``. This
legacy V2 renderer now refuses to treat a short delta as a complete manuscript
unless an abridged derivative is explicitly requested.
"""
import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 55, 72)
MUTED = RGBColor(92, 99, 108)
GOLD = RGBColor(150, 111, 32)
LIGHT = "F4F6F9"

EXPORTS = [
    {
        "project": REPO / "examples/sri_lanka_pre_1948",
        "output": "Sri_Lanka_Fresque_historico_geographique_vol_retour_v2.docx",
        "kicker": "FRESQUE HISTORICO-GÉOGRAPHIQUE · VOLUME 1",
        "title": "Sri Lanka — des interfaces anciennes à l’État de 1948",
        "subtitle": "Jaffna, Polonnaruwa, réseaux de l’océan Indien, empires côtiers et héritages administratifs",
        "running": "Sri Lanka · longue durée jusqu’à 1948",
    },
    {
        "project": REPO / "examples/sri_lanka_post_1948",
        "output": "Sri_Lanka_1948_2026_etude_historico_geographique_v2.docx",
        "kicker": "FRESQUE HISTORICO-GÉOGRAPHIQUE · VOLUME 2",
        "title": "Sri Lanka 1948–2026",
        "subtitle": "État, langue, caste, guerre, diaspora et conversion territoriale du capital humain",
        "running": "Sri Lanka · 1948–2026",
    },
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = edges.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_section(section, running):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(running)
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_field(footer.paragraphs[0])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Causal Callout" not in doc.styles:
        callout = doc.styles.add_style("Causal Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Causal Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(11)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.208
    callout.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_paragraph_fill(paragraph, fill=LIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2E74B5")
    left.set(qn("w:space"), "8")
    borders.append(left)
    p_pr.append(borders)


def add_inline_markdown(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, italic=True)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]))


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def cover(doc, spec):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(118)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(spec["kicker"]), size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(spec["title"]), size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    set_run_font(p.add_run(spec["subtitle"]), size=15, color=DARK_BLUE)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    for cell in table.rows[0].cells:
        shade_cell(cell, "F4F6F9")
        set_cell_border(cell, top={"val":"single","sz":"4","color":"D7DBE2"}, bottom={"val":"single","sz":"4","color":"D7DBE2"})
    left, right = table.rows[0].cells
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(left_p.add_run("ÉDITION LECTEUR V2\n"), size=9, color=MUTED, bold=True)
    set_run_font(left_p.add_run("Markdown canonique Run 5 · Storytelling Run 6"), size=9.5, color=INK)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right_p.add_run("18 AOÛT 2026\n"), size=9, color=MUTED, bold=True)
    set_run_font(right_p.add_run("Version vérifiée et traçable"), size=9.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(62)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Une lecture chronologique guidée par les lieux, les mécanismes et leurs sources."), size=10, color=MUTED, italic=True)
    doc.add_page_break()


def collect_register(project):
    items = {}
    for path in sorted((project / "05_sources").glob("source_register*.json")):
        for item in json.loads(path.read_text(encoding="utf-8")):
            items[item["id"]] = item
    return items


def extract_source_ids(text, register):
    used = []
    for source_id in register:
        if source_id in text:
            used.append(source_id)
    return sorted(used)


def add_reading_map(doc, headings):
    doc.add_heading("Repères de lecture", level=1)
    p = doc.add_paragraph()
    add_inline_markdown(p, "Cette édition conserve la chronologie comme colonne vertébrale. Les encadrés signalent un mécanisme, une fausse piste ou un changement d’échelle ; les identifiants entre crochets renvoient au registre de sources du projet.")
    midpoint = (len(headings) + 1) // 2
    columns = [headings[:midpoint], headings[midpoint:]]
    rows = max(len(column) for column in columns)
    table = doc.add_table(rows=rows, cols=2)
    set_table_geometry(table, [4680, 4680])
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_border(cell, bottom={"val":"single","sz":"3","color":"E1E6EC"})
            item_index = row_index
            if item_index >= len(columns[column_index]):
                continue
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            number = row_index + 1 + (midpoint if column_index else 0)
            set_run_font(p.add_run(f"{number:02d}  "), size=9.5, color=GOLD, bold=True)
            add_inline_markdown(p, re.sub(r"^\d+\.\s*", "", columns[column_index][item_index]))
    doc.add_page_break()


def render_markdown(doc, text):
    lines = text.splitlines()
    body_started = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            body_started = True
            continue
        if stripped.startswith("> **Statut**"):
            continue
        if stripped == "---":
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline_markdown(p, stripped[4:])
        elif stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline_markdown(p, stripped[3:])
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Causal Callout")
            add_paragraph_fill(p)
            add_inline_markdown(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\. ", "", stripped))
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
        elif body_started:
            p = doc.add_paragraph()
            add_inline_markdown(p, stripped)


def add_sources_appendix(doc, source_ids, register):
    doc.add_page_break()
    doc.add_heading("Sources citées dans cette édition", level=1)
    p = doc.add_paragraph()
    add_inline_markdown(p, "Les tiers décrivent le mode de production de la connaissance ; le rôle décrit l’usage dans cette enquête. Les limites sont conservées pour empêcher qu’une source soit étendue au-delà de son champ.")
    for source_id in source_ids:
        source = register[source_id]
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{source_id} — {source.get('title', '')}")
        set_run_font(run, bold=True, color=INK)
        meta = doc.add_paragraph()
        meta.paragraph_format.left_indent = Inches(0.2)
        meta.paragraph_format.space_after = Pt(4)
        add_inline_markdown(meta, f"{source.get('tier', '')} · {source.get('anchor_role', '')}. {source.get('scope', '')}. Limite : {source.get('limitations', '')}. ")
        add_hyperlink(meta, "Lien source", source.get("url", ""))


def audit_preset(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    assert all(round(v.inches, 3) == 1.0 for v in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin))
    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri" and round(normal.font.size.pt, 1) == 11.0
    assert round(normal.paragraph_format.space_after.pt, 1) == 8.0
    assert round(float(normal.paragraph_format.line_spacing), 3) == 1.333
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        assert round(doc.styles[name].font.size.pt, 1) == size


def enforce_advanced_retention(project, text, allow_abridged=False):
    baseline = project / "09_output" / "report_v1_full.md"
    if not baseline.exists() or allow_abridged:
        return
    candidate_words = len(text.split())
    baseline_words = len(baseline.read_text(encoding="utf-8").split())
    if candidate_words < baseline_words:
        raise RuntimeError(
            "Refusing silent advanced-reader compression: report.md contains "
            f"{candidate_words} words but the complete V1 baseline contains "
            f"{baseline_words}. Use render_full_reader_v3.py, or explicitly pass "
            "--allow-abridged for a labelled derivative."
        )


def build(spec, allow_abridged=False):
    project = spec["project"]
    markdown = project / "09_output" / "report.md"
    text = markdown.read_text(encoding="utf-8")
    enforce_advanced_retention(project, text, allow_abridged=allow_abridged)
    headings = [line[3:].strip() for line in text.splitlines() if line.startswith("## ")]
    register = collect_register(project)
    source_ids = extract_source_ids(text, register)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], spec["running"])
    cover(doc, spec)
    add_reading_map(doc, headings)
    render_markdown(doc, text)
    add_sources_appendix(doc, source_ids, register)

    core = doc.core_properties
    core.title = spec["title"]
    core.subject = "Fresque historico-géographique — édition lecteur v2"
    core.author = "Projet tourisme-etude-historico-geographique"
    core.keywords = "Sri Lanka, histoire, géographie, Jaffna, Polonnaruwa, reader edition"
    core.comments = "Generated from promoted Markdown on 2026-08-18; narrative_proposal preset with editorial_cover."

    audit_preset(doc)
    destination = project / "09_output" / spec["output"]
    doc.save(destination)
    print(destination.relative_to(REPO))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--project", choices=["pre", "post", "all"], default="all")
    parser.add_argument("--allow-abridged", action="store_true")
    args = parser.parse_args()
    choices = EXPORTS if args.project == "all" else [EXPORTS[0 if args.project == "pre" else 1]]
    for spec in choices:
        build(spec, allow_abridged=args.allow_abridged)


if __name__ == "__main__":
    main()
