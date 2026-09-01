#!/usr/bin/env python3
"""Place reader-eligible side stories inside reviewed host paragraphs.

Run41 contract:
- default placement is inside the logical host paragraph at a safe sentence boundary;
- at most one embedded side story per host paragraph;
- in any three consecutive logical host paragraphs, at most two remain embedded;
  the story with the best interstitial affinity/transition fit becomes a local
  inter-paragraph excursion;
- at most one interstitial story per paragraph boundary;
- no append-at-book-end fallback and no sentence splitting.
"""
from __future__ import annotations

import json
import re
import string
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from materialize_side_stories import (
    side_story_begin_marker,
    side_story_end_marker,
    validate_narrative_depth,
)
from side_story_contract import canonical_marker, load_side_stories
from side_story_presentation import style_side_story_cell

ELIGIBILITY_BASIS = "museum_plus_independent_corroboration"
PLAN_NAME = "side_story_placement_plan.json"
SIDE_MARKER = "[SIDE-STORY:"
PROFILE_PATH = Path(__file__).resolve().parents[1] / "templates" / "side-stories" / "type_profiles.json"


def _profiles() -> dict:
    return json.loads(PROFILE_PATH.read_text(encoding="utf-8"))


def _profile(kind: str) -> dict:
    profile = (_profiles().get("profiles") or {}).get(kind)
    if not isinstance(profile, dict):
        raise RuntimeError(f"missing side-story type profile for {kind!r}")
    return profile


def _placement_contract() -> dict:
    return _profiles().get("placement_contract") or {}


def _norm(value: str) -> str:
    value = re.sub(r"<[^>]+>", " ", str(value))
    value = re.sub(r"[#*_`>\[\](){}]", " ", value)
    value = value.translate(
        str.maketrans({c: " " for c in string.punctuation + "«»“”‘’…–—≠×"})
    )
    return re.sub(r"\s+", " ", value).strip().casefold()


def _visible_words(value: str) -> int:
    value = re.sub(r"<[^>]+>", " ", value)
    value = re.sub(r"\[[^\]]+\]\([^\)]+\)", r"\1", value)
    return len(
        re.findall(r"\b[\wÀ-ÖØ-öø-ÿĀ-ž'’.-]+\b", value, re.UNICODE)
    )


def _plain(value: str) -> str:
    value = re.sub(r"<!--.*?-->", " ", value, flags=re.S)
    value = re.sub(r"\[[^\]]+\]\([^\)]+\)", r"\1", value)
    value = re.sub(r"^#{1,6}\s+", "", value.strip())
    value = re.sub(r"[*_`]", "", value)
    return re.sub(r"\s+", " ", value).strip()


def _content_body(item: dict) -> str:
    return ((item.get("content") or {}).get("body_markdown") or "").strip()


def candidate_reader_eligible(item: dict) -> tuple[bool, list[str]]:
    """Return reader eligibility without changing candidate evidence status."""
    if item.get("status") != "candidate":
        return False, []
    q = item.get("reader_eligibility") or {}
    errors: list[str] = []
    if q.get("basis") != ELIGIBILITY_BASIS:
        return False, []
    if q.get("status") != "eligible":
        errors.append("reader_eligibility.status must be eligible")
    if q.get("forced_pipeline") is not True:
        errors.append("reader_eligibility.forced_pipeline must be true")
    museum = list(q.get("museum_source_ids") or [])
    corroborating = list(q.get("corroborating_source_ids") or [])
    origin = list(q.get("museum_origin_paths") or [])
    if not museum and not origin:
        errors.append("museum source or museum origin path required")
    if not corroborating:
        errors.append("at least one independent corroborating source required")
    if set(museum) & set(corroborating):
        errors.append("museum and corroborating source sets must be disjoint")
    lineage_sources = set((item.get("lineage") or {}).get("source_ids") or [])
    if set(museum) - lineage_sources:
        errors.append("museum_source_ids must be in lineage.source_ids")
    if set(corroborating) - lineage_sources:
        errors.append("corroborating_source_ids must be in lineage.source_ids")
    if not q.get("uncertainty_boundary"):
        errors.append("explicit uncertainty_boundary required")
    if not _content_body(item):
        errors.append("reader-eligible candidate requires content.body_markdown")
    placement = item.get("placement") or {}
    if not placement.get("section_anchor"):
        errors.append("reader-eligible candidate requires placement.section_anchor")
    if not (placement.get("match_terms") or []):
        errors.append("reader-eligible candidate requires placement.match_terms")
    return not errors, errors


def is_post_review_materializable(item: dict) -> bool:
    if not (item.get("render") or {}).get("required_in_reader"):
        return False
    if item.get("status") in {"promoted", "validated"}:
        return True
    eligible, errors = candidate_reader_eligible(item)
    if errors:
        raise RuntimeError(f"{item.get('id')}: " + "; ".join(errors))
    return eligible


def _markdown_story_block(item: dict) -> str:
    body = _content_body(item)
    if not body:
        raise RuntimeError(
            f"{item.get('id')}: post-review placement requires body_markdown"
        )
    validate_narrative_depth(item, body)
    label = (item.get("render") or {}).get("label") or item.get("kind")
    prefix = f"**{label} — {item['title']}**"
    return "\n".join(
        [side_story_begin_marker(item), prefix, "", body, side_story_end_marker(item)]
    )


def _split_blocks(text: str) -> list[str]:
    return re.split(r"\n{2,}", text.strip())


def _is_plain_host(block: str) -> bool:
    s = block.strip()
    if not s or SIDE_MARKER in s:
        return False
    if s.startswith(("#", "<!--", "```", "|", "<table", "<tr", "<td", "---", ">")):
        return False
    if re.match(r"^(?:[-*+] |\d+[.)] )", s):
        return False
    return _visible_words(s) >= 12


def _heading_level(block: str) -> int | None:
    first = block.strip().splitlines()[0] if block.strip() else ""
    m = re.match(r"^(#{1,6})\s+", first)
    return len(m.group(1)) if m else None


def _looks_like_reader_heading(block: str, anchor: str) -> bool:
    s = block.strip()
    if not s or s.startswith(">"):
        return False
    first = s.splitlines()[0]
    if _heading_level(block) is not None:
        return True
    plain = re.sub(r"[*_`]", "", first).strip()
    return (
        _norm(anchor) in _norm(plain)
        and _visible_words(plain) <= 24
        and len(s.splitlines()) <= 2
    )


def _find_section_range(blocks: list[str], anchor: str) -> tuple[int, int]:
    needle = _norm(anchor)
    hits = [i for i, b in enumerate(blocks) if needle and needle in _norm(b)]
    if not hits:
        raise RuntimeError(f"section anchor not found in reviewed reader: {anchor}")
    explicit = [i for i in hits if _looks_like_reader_heading(blocks[i], anchor)]
    hit = explicit[-1] if explicit else hits[-1]
    start = hit
    level = _heading_level(blocks[start])
    if level is None:
        heading = next(
            (i for i in range(start, -1, -1) if _heading_level(blocks[i]) is not None),
            None,
        )
        if heading is None:
            end = next(
                (
                    i
                    for i in range(start + 1, len(blocks))
                    if _looks_like_reader_heading(
                        blocks[i],
                        re.sub(r"[*_`]", "", blocks[i].splitlines()[0]),
                    )
                ),
                len(blocks),
            )
            return start, end
        start = heading
        level = _heading_level(blocks[start])
    end = len(blocks)
    for i in range(start + 1, len(blocks)):
        lv = _heading_level(blocks[i])
        if lv is not None and lv <= level:
            end = i
            break
    return start, end


def _host_ordinals(blocks: list[str]) -> dict[int, int]:
    return {idx: ordinal for ordinal, idx in enumerate(i for i, b in enumerate(blocks) if _is_plain_host(b))}


def _existing_story_near(blocks: list[str], idx: int) -> int:
    return sum(
        1
        for j in range(max(0, idx - 1), min(len(blocks), idx + 2))
        if SIDE_MARKER in blocks[j]
    )


def _score_host(blocks: list[str], idx: int, item: dict) -> dict:
    placement = item.get("placement") or {}
    norm = _norm(blocks[idx])
    terms = [_norm(x) for x in placement.get("match_terms") or [] if _norm(x)]
    chronology = [
        _norm(x) for x in placement.get("chronology_terms") or [] if _norm(x)
    ]
    mechanism = [
        _norm(x) for x in placement.get("mechanism_terms") or [] if _norm(x)
    ]
    term_hits = [t for t in terms if t in norm]
    chronology_hits = [t for t in chronology if t in norm]
    mechanism_hits = [t for t in mechanism if t in norm]
    semantic = (
        120 * len(term_hits)
        + 80 * len(chronology_hits)
        + 70 * len(mechanism_hits)
    )
    existing_penalty = 220 * _existing_story_near(blocks, idx)
    words = _visible_words(blocks[idx])
    richness = min(words, 180) / 18.0
    return {
        "score": semantic - existing_penalty + richness,
        "semantic": semantic,
        "existing_density_penalty": existing_penalty,
        "words": words,
        "term_hits": term_hits,
        "chronology_hits": chronology_hits,
        "mechanism_hits": mechanism_hits,
    }


def _rank_hosts(blocks: list[str], item: dict) -> list[tuple[int, dict]]:
    placement = item.get("placement") or {}
    start, end = _find_section_range(blocks, placement["section_anchor"])
    candidates = [
        (i, _score_host(blocks, i, item))
        for i in range(start + 1, end)
        if _is_plain_host(blocks[i])
    ]
    if not candidates:
        raise RuntimeError(
            f"{item.get('id')}: no host paragraph inside section {placement['section_anchor']!r}"
        )
    semantic = [row for row in candidates if row[1]["semantic"] > 0]
    fallback = not semantic
    pool = semantic or candidates
    if fallback:
        pool = sorted(
            pool,
            key=lambda row: (
                -_existing_story_near(blocks, row[0]),
                row[1]["words"],
                row[1]["score"],
            ),
            reverse=True,
        )
    else:
        pool = sorted(
            pool, key=lambda row: (row[1]["score"], row[1]["words"]), reverse=True
        )
    out: list[tuple[int, dict]] = []
    for idx, meta in pool:
        m = dict(meta)
        m["fallback_longest_fragment"] = fallback
        m["section_range"] = [start, end]
        out.append((idx, m))
    return out


def _sentence_breaks(text: str) -> list[int]:
    """Candidate raw-string offsets *after* terminal punctuation."""
    breaks: list[int] = []
    for m in re.finditer(r"[.!?…](?=\s|$)", text):
        pos = m.end()
        if _visible_words(text[:pos]) >= 12 and _visible_words(text[pos:]) >= 12:
            breaks.append(pos)
    return breaks


def _choose_sentence_split(block: str, item: dict) -> tuple[int, int] | None:
    breaks = _sentence_breaks(block)
    if not breaks:
        return None
    placement = item.get("placement") or {}
    terms = [
        _norm(x)
        for x in (
            list(placement.get("match_terms") or [])
            + list(placement.get("mechanism_terms") or [])
            + list(placement.get("chronology_terms") or [])
        )
        if _norm(x)
    ]
    target_ratio = 0.55
    total = max(_visible_words(block), 1)
    scored: list[tuple[float, int, int]] = []
    previous = 0
    for pos in breaks:
        sentence = block[previous:pos]
        hits = sum(1 for t in terms if t and t in _norm(sentence))
        before_words = _visible_words(block[:pos])
        ratio = before_words / total
        score = 100.0 * hits - abs(ratio - target_ratio) * 10.0
        scored.append((score, pos, before_words))
        previous = pos
    _, pos, before_words = max(scored, key=lambda row: row[0])
    return pos, before_words


def _interstitial_affinity(item: dict) -> int:
    return int(_profile(str(item.get("kind") or "")).get("interstitial_affinity", 0))


def _transition_fit(item: dict, blocks: list[str], host_idx: int) -> int:
    """How well the story can stand between the host and its next/previous paragraph."""
    placement = item.get("placement") or {}
    terms = [
        _norm(x)
        for x in (
            list(placement.get("match_terms") or [])
            + list(placement.get("mechanism_terms") or [])
        )
        if _norm(x)
    ]
    score = 0
    for j in (host_idx - 1, host_idx + 1):
        if 0 <= j < len(blocks) and _is_plain_host(blocks[j]):
            n = _norm(blocks[j])
            score += sum(1 for t in terms if t and t in n)
    return score


def _assign_initial(blocks: list[str], queue: list[dict]) -> list[dict]:
    host_ord = _host_ordinals(blocks)
    occupied: set[int] = set()
    assignments: list[dict] = []
    for item in queue:
        ranked = _rank_hosts(blocks, item)
        chosen_idx = None
        chosen_meta = None
        split = None
        for idx, meta in ranked:
            candidate_split = _choose_sentence_split(blocks[idx], item)
            if idx not in occupied and candidate_split is not None:
                chosen_idx, chosen_meta, split = idx, meta, candidate_split
                break
        if chosen_idx is None:
            chosen_idx, chosen_meta = ranked[0]
            mode = "interstitial"
            split_words = None
            split_raw = None
        else:
            occupied.add(chosen_idx)
            mode = "embedded"
            split_raw, split_words = split
        assignments.append(
            {
                "item": item,
                "host_idx": chosen_idx,
                "host_ordinal": host_ord[chosen_idx],
                "meta": chosen_meta,
                "placement_mode": mode,
                "split_raw_offset": split_raw,
                "split_after_visible_words": split_words,
                "density_reason": None,
            }
        )
    return assignments


def _enforce_three_paragraph_density(blocks: list[str], assignments: list[dict]) -> None:
    """Convert one of three consecutive embedded stories to interstitial."""
    changed = True
    while changed:
        changed = False
        embedded = {
            a["host_ordinal"]: a
            for a in assignments
            if a["placement_mode"] == "embedded"
        }
        ordinals = sorted(embedded)
        if not ordinals:
            return
        low, high = min(ordinals), max(ordinals)
        for start in range(low, high - 1):
            triple = [start, start + 1, start + 2]
            if not all(o in embedded for o in triple):
                continue
            candidates = [embedded[o] for o in triple]
            chosen = max(
                candidates,
                key=lambda a: (
                    _interstitial_affinity(a["item"]),
                    _transition_fit(a["item"], blocks, a["host_idx"]),
                    -abs(a["host_ordinal"] - (start + 1)),
                ),
            )
            chosen["placement_mode"] = "interstitial"
            chosen["split_raw_offset"] = None
            chosen["split_after_visible_words"] = None
            chosen["density_reason"] = "three_paragraph_window_overflow"
            changed = True
            break


def _reserve_interstitial_boundaries(assignments: list[dict]) -> None:
    """At most one story per logical paragraph boundary."""
    reserved: set[tuple[int, str]] = set()
    for a in assignments:
        if a["placement_mode"] != "interstitial":
            continue
        pref = (a["item"].get("placement") or {}).get("position", "after")
        preferred = "before" if pref == "before" else "after"
        candidates = [
            (a["host_ordinal"], preferred),
            (a["host_ordinal"], "before" if preferred == "after" else "after"),
            (a["host_ordinal"] + 1, "before"),
            (max(0, a["host_ordinal"] - 1), "after"),
        ]
        boundary = next((b for b in candidates if b not in reserved), None)
        if boundary is None:
            raise RuntimeError(
                f"{a['item'].get('id')}: no free local interstitial boundary"
            )
        reserved.add(boundary)
        a["interstitial_boundary"] = {
            "host_ordinal": boundary[0],
            "position": boundary[1],
        }


def _plan_assignments(blocks: list[str], queue: list[dict]) -> list[dict]:
    assignments = _assign_initial(blocks, queue)
    _enforce_three_paragraph_density(blocks, assignments)
    _reserve_interstitial_boundaries(assignments)
    return assignments


def _render_markdown(blocks: list[str], assignments: list[dict]) -> tuple[str, list[dict]]:
    embedded_by_idx = {
        a["host_idx"]: a
        for a in assignments
        if a["placement_mode"] == "embedded"
    }
    before: dict[int, list[dict]] = {}
    after: dict[int, list[dict]] = {}
    ord_to_idx = {v: k for k, v in _host_ordinals(blocks).items()}
    for a in assignments:
        if a["placement_mode"] != "interstitial":
            continue
        b = a["interstitial_boundary"]
        idx = ord_to_idx.get(b["host_ordinal"])
        if idx is None:
            raise RuntimeError(
                f"{a['item'].get('id')}: interstitial boundary host ordinal missing"
            )
        (before if b["position"] == "before" else after).setdefault(idx, []).append(a)

    out: list[str] = []
    ledger: list[dict] = []
    for idx, block in enumerate(blocks):
        for a in before.get(idx, []):
            out.append(_markdown_story_block(a["item"]))
        if idx in embedded_by_idx:
            a = embedded_by_idx[idx]
            split = int(a["split_raw_offset"])
            left = block[:split].rstrip()
            right = block[split:].lstrip()
            if not left or not right:
                raise RuntimeError(
                    f"{a['item'].get('id')}: invalid empty host segment after split"
                )
            out.extend([left, _markdown_story_block(a["item"]), right])
        else:
            out.append(block)
        for a in after.get(idx, []):
            out.append(_markdown_story_block(a["item"]))

    for a in assignments:
        meta = a["meta"]
        row = {
            "id": a["item"]["id"],
            "kind": a["item"].get("kind"),
            "status": a["item"].get("status"),
            "section_anchor": (a["item"].get("placement") or {}).get("section_anchor"),
            "placement_mode": a["placement_mode"],
            "host_paragraph_ordinal": a["host_ordinal"],
            "host_paragraph_excerpt": _plain(blocks[a["host_idx"]])[:260],
            "semantic_score": meta["semantic"],
            "score": round(meta["score"], 2),
            "existing_density_penalty": meta["existing_density_penalty"],
            "paragraph_words": meta["words"],
            "term_hits": meta["term_hits"],
            "chronology_hits": meta["chronology_hits"],
            "mechanism_hits": meta["mechanism_hits"],
            "fallback_longest_fragment": meta["fallback_longest_fragment"],
            "split_after_visible_words": a.get("split_after_visible_words"),
            "density_reason": a.get("density_reason"),
            "interstitial_affinity": _interstitial_affinity(a["item"]),
        }
        if a["placement_mode"] == "interstitial":
            row["interstitial_boundary"] = a["interstitial_boundary"]
            boundary_idx = ord_to_idx[a["interstitial_boundary"]["host_ordinal"]]
            row["boundary_host_excerpt"] = _plain(blocks[boundary_idx])[:260]
        ledger.append(row)
    return "\n\n".join(out).strip() + "\n", ledger


def place_markdown(project: Path, text: str) -> tuple[str, list[dict]]:
    blocks = _split_blocks(text)
    queue: list[dict] = []
    for _, item in load_side_stories(project):
        if not is_post_review_materializable(item):
            continue
        marker = canonical_marker(item["id"])
        if marker in text or item.get("materialization_mode") == "existing_fragment":
            continue
        queue.append(item)
    queue.sort(
        key=lambda item: _find_section_range(
            blocks, (item.get("placement") or {})["section_anchor"]
        )[0]
    )
    assignments = _plan_assignments(blocks, queue)
    return _render_markdown(blocks, assignments)


def _find_docx_paragraph(doc: Document, excerpt: str) -> Paragraph | None:
    needle = _norm(excerpt)
    if not needle:
        return None
    for p in doc.paragraphs:
        n = _norm(p.text)
        if n and (needle in n or n in needle) and min(len(n), len(needle)) >= 35:
            return p
    short = " ".join(needle.split()[:12])
    for p in doc.paragraphs:
        if short and short in _norm(p.text):
            return p
    return None


def _set_para_text_preserve_style(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _split_after_visible_words(text: str, count: int) -> tuple[str, str]:
    words = list(re.finditer(r"\b[\wÀ-ÖØ-öø-ÿĀ-ž'’.-]+\b", text, re.UNICODE))
    if count <= 0 or count >= len(words):
        raise RuntimeError("invalid DOCX word split")
    pos = words[count - 1].end()
    tail = text[pos : pos + 8]
    punct = re.match(r"^[\s]*[.!?…]", tail)
    if punct:
        pos += punct.end()
    return text[:pos].rstrip(), text[pos:].lstrip()


def _insert_story_table_after(doc: Document, paragraph: Paragraph, item: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    style_side_story_cell(cell, item["kind"])
    label = (item.get("render") or {}).get("label") or item.get("kind")
    head = cell.paragraphs[0]
    run = head.add_run(f"{label} — {item['title']}")
    run.bold = True
    for block in re.split(r"\n{2,}", _content_body(item)):
        text = _plain(block)
        if text:
            cell.add_paragraph(text)
    paragraph._p.addnext(table._tbl)


def _insert_story_table_before(doc: Document, paragraph: Paragraph, item: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    style_side_story_cell(cell, item["kind"])
    label = (item.get("render") or {}).get("label") or item.get("kind")
    head = cell.paragraphs[0]
    run = head.add_run(f"{label} — {item['title']}")
    run.bold = True
    for block in re.split(r"\n{2,}", _content_body(item)):
        text = _plain(block)
        if text:
            cell.add_paragraph(text)
    paragraph._p.addprevious(table._tbl)


def _split_docx_and_embed(doc: Document, paragraph: Paragraph, item: dict, split_words: int) -> None:
    full = paragraph.text
    left, right = _split_after_visible_words(full, split_words)
    if _visible_words(left) < 8 or _visible_words(right) < 8:
        raise RuntimeError(
            f"{item.get('id')}: DOCX host split leaves insufficient context"
        )
    right_xml = deepcopy(paragraph._p)
    paragraph._p.addnext(right_xml)
    right_p = Paragraph(right_xml, paragraph._parent)
    _set_para_text_preserve_style(paragraph, left)
    _set_para_text_preserve_style(right_p, right)
    _insert_story_table_after(doc, paragraph, item)


def place_docx(project: Path, docx_path: Path, plan: list[dict]) -> int:
    if not plan:
        return 0
    by_id = {item["id"]: item for _, item in load_side_stories(project)}
    doc = Document(docx_path)
    inserted = 0

    for row in [r for r in plan if r["placement_mode"] == "interstitial"]:
        item = by_id[row["id"]]
        paragraph = _find_docx_paragraph(doc, row.get("boundary_host_excerpt") or row["host_paragraph_excerpt"])
        if paragraph is None:
            raise RuntimeError(
                f"{row['id']}: interstitial host could not be resolved in DOCX core"
            )
        boundary = row["interstitial_boundary"]
        if boundary["position"] == "before":
            _insert_story_table_before(doc, paragraph, item)
        else:
            _insert_story_table_after(doc, paragraph, item)
        inserted += 1

    for row in [r for r in plan if r["placement_mode"] == "embedded"]:
        item = by_id[row["id"]]
        paragraph = _find_docx_paragraph(doc, row["host_paragraph_excerpt"])
        if paragraph is None:
            raise RuntimeError(
                f"{row['id']}: embedded host could not be resolved in DOCX core"
            )
        _split_docx_and_embed(
            doc, paragraph, item, int(row["split_after_visible_words"])
        )
        inserted += 1

    doc.save(docx_path)
    return inserted


def place_reader(project: Path, markdown_path: Path, docx_path: Path) -> dict:
    core = markdown_path.read_text(encoding="utf-8")
    placed, plan = place_markdown(project, core)
    markdown_path.write_text(placed, encoding="utf-8")
    docx_count = place_docx(project, docx_path, plan)
    contract = _placement_contract()
    plan_path = project / "09_output" / PLAN_NAME
    plan_path.write_text(
        json.dumps(
            {
                "schema_version": "1.1",
                "stage": "post_core_review",
                "placement_contract": contract,
                "placements": plan,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    return {
        "post_review_side_stories": len(plan),
        "post_review_side_stories_docx": docx_count,
        "side_story_embedded": sum(
            1 for row in plan if row["placement_mode"] == "embedded"
        ),
        "side_story_interstitial": sum(
            1 for row in plan if row["placement_mode"] == "interstitial"
        ),
        "side_story_placement_plan": str(plan_path),
    }
