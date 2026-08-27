#!/usr/bin/env python3
"""Render a reviewed Run26 from-scratch Markdown manuscript into a fresh DOCX.

Unlike render_full_reader_v3.py this renderer never opens a previous reader DOCX.
Every side-story BEGIN/END pair becomes one pastel, bordered one-cell container.
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

from frontstage_reader_contract import assert_no_known_backstage_leak,visible_docx_text
from from_scratch_review_contract import CLAIM_MARKER,SIDE_BEGIN,SIDE_END,assert_review_complete
from render_reader_exports import add_inline_markdown,configure_section,configure_styles,set_run_font
from side_story_contract import RENDER_LABELS
from side_story_presentation import SIDE_STORY_PRESENTATION,add_side_story_legend,style_side_story_cell

REPO=Path(__file__).resolve().parents[1]
INK=RGBColor(32,55,72);MUTED=RGBColor(92,99,108);GOLD=RGBColor(150,111,32)

SPECS={
    "pre":{
        "project":REPO/"examples/sri_lanka_pre_1948",
        "manuscript":"Sri_Lanka_pre_1948_run26_from_scratch.md",
        "ledger":"review_ledger_pre.json",
        "output":"Sri_Lanka_pre_1948_run26_from_scratch.docx",
        "title":"Sri Lanka — une île construite par ses interfaces",
        "subtitle":"Des premiers réseaux de l’océan Indien à l’État colonial de 1948",
        "running":"Sri Lanka · des origines à 1948",
    },
    "post":{
        "project":REPO/"examples/sri_lanka_post_1948",
        "manuscript":"Sri_Lanka_post_1948_run26_from_scratch.md",
        "ledger":"review_ledger_post.json",
        "output":"Sri_Lanka_post_1948_run26_from_scratch.docx",
        "title":"Sri Lanka depuis 1948 — l’État, les territoires et leurs fractures",
        "subtitle":"Langue, mobilité, guerre, environnement, plantations et recompositions contemporaines",
        "running":"Sri Lanka · 1948–2026",
    },
}


def clean_reader_text(text:str)->str:
    return re.sub(r"\s+"," ",CLAIM_MARKER.sub("",text)).strip()


def set_cell_width_auto(table):
    table.autofit=True
    tbl_pr=table._tbl.tblPr
    tbl_w=tbl_pr.find(qn("w:tblW"))
    if tbl_w is not None:tbl_w.set(qn("w:type"),"auto");tbl_w.set(qn("w:w"),"0")


def add_cover(doc:Document,spec:dict)->None:
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(110);p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("FRESQUE HISTORICO-GÉOGRAPHIQUE"),size=10,color=GOLD,bold=True)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(16);p.paragraph_format.space_after=Pt(10)
    set_run_font(p.add_run(spec["title"]),size=28,color=INK,bold=True)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(64)
    set_run_font(p.add_run(spec["subtitle"]),size=14,color=MUTED)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Une narration reconstruite depuis les artefacts historiques structurés."),size=10,color=MUTED,italic=True)
    doc.add_page_break()


def add_toc_field(doc:Document)->None:
    heading=doc.add_heading("Sommaire",level=1);heading.paragraph_format.keep_with_next=True
    p=doc.add_paragraph()
    begin=OxmlElement("w:fldChar");begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText");instr.set(qn("xml:space"),"preserve");instr.text=' TOC \\o "1-3" \\h \\z \\u '
    separate=OxmlElement("w:fldChar");separate.set(qn("w:fldCharType"),"separate")
    text=OxmlElement("w:t");text.text="Mettre à jour le sommaire dans Word si nécessaire."
    run=OxmlElement("w:r");run.append(text)
    end=OxmlElement("w:fldChar");end.set(qn("w:fldCharType"),"end")
    p._p.extend([begin,instr,separate,run,end])
    settings=doc.settings._element;update=settings.find(qn("w:updateFields"))
    if update is None:update=OxmlElement("w:updateFields");settings.append(update)
    update.set(qn("w:val"),"true")
    doc.add_page_break()


def _paragraph(target,style=None):
    return target.add_paragraph(style=style) if style else target.add_paragraph()


def add_line(target,line:str,inside_side_story:bool=False)->None:
    text=clean_reader_text(line.strip())
    if not text:return
    heading=re.match(r"^(#{1,4})\s+(.+)$",text)
    if heading and not inside_side_story:
        level=min(len(heading.group(1)),3);p=target.add_heading(level=level);add_inline_markdown(p,heading.group(2));return
    if heading and inside_side_story:
        p=_paragraph(target);run=p.add_run(re.sub(r"^#{1,4}\s+","",text));run.bold=True;return
    if re.match(r"^\d+\.\s+",text):
        p=_paragraph(target,"List Number");add_inline_markdown(p,re.sub(r"^\d+\.\s+","",text));return
    if text.startswith(("- ","* ")):
        p=_paragraph(target,"List Bullet");add_inline_markdown(p,text[2:]);return
    if text.startswith("> "):
        p=_paragraph(target);p.paragraph_format.left_indent=Pt(12);add_inline_markdown(p,text[2:]);return
    p=_paragraph(target);add_inline_markdown(p,text)


def add_side_story_block(doc:Document,kind:str,body_lines:list[str])->None:
    if kind not in SIDE_STORY_PRESENTATION:raise RuntimeError(f"unknown side-story kind in manuscript: {kind}")
    table=doc.add_table(rows=1,cols=1);set_cell_width_auto(table);cell=table.cell(0,0);style_side_story_cell(cell,kind)
    # Remove default empty text only by reusing its first paragraph.
    cell.paragraphs[0].text=""
    first_nonempty=True
    for raw in body_lines:
        s=raw.strip()
        if not s:continue
        text=clean_reader_text(s)
        if first_nonempty:
            expected=RENDER_LABELS[kind]
            if expected.casefold() not in re.sub(r"[*_]","",text).casefold():
                raise RuntimeError(f"side-story block kind={kind} does not start with its reader label {expected!r}")
            symbol=SIDE_STORY_PRESENTATION[kind]["symbol"]
            text=re.sub(r"^\*\*",f"**{symbol} ",text,count=1) if text.startswith("**") else f"{symbol} {text}"
            p=cell.paragraphs[0];add_inline_markdown(p,text);p.paragraph_format.keep_with_next=True;first_nonempty=False;continue
        add_line(cell,text,inside_side_story=True)
    if first_nonempty:raise RuntimeError(f"empty side-story block kind={kind}")
    doc.add_paragraph().paragraph_format.space_after=Pt(2)


def render_markdown(doc:Document,markdown:str)->dict:
    side_story_count=0;normal_lines=[];inside=None;body=[]
    def flush_normal():
        nonlocal normal_lines
        for raw in normal_lines:add_line(doc,raw,False)
        normal_lines=[]
    for raw in markdown.splitlines():
        stripped=raw.strip();begin=SIDE_BEGIN.match(stripped);end=SIDE_END.match(stripped)
        if begin:
            if inside is not None:raise RuntimeError("nested side-story blocks are not allowed")
            flush_normal();inside={"id":begin.group(1),"kind":begin.group(2)};body=[];continue
        if end:
            if inside is None:raise RuntimeError(f"side-story END without BEGIN: {end.group(1)}")
            if end.group(1)!=inside["id"]:raise RuntimeError(f"side-story fence mismatch: {inside['id']} != {end.group(1)}")
            add_side_story_block(doc,inside["kind"],body);side_story_count+=1;inside=None;body=[];continue
        if stripped.startswith("<!--"):continue
        if inside is not None:body.append(raw)
        else:normal_lines.append(raw)
    if inside is not None:raise RuntimeError(f"unclosed side-story block: {inside['id']}")
    flush_normal();return {"side_story_blocks":side_story_count}


def build(key:str)->dict:
    spec=SPECS[key];project=spec["project"];root=project/"09_output"/"from_scratch";manuscript=root/spec["manuscript"];ledger=root/spec["ledger"]
    reviewed=assert_review_complete(project,manuscript,ledger)
    markdown=manuscript.read_text(encoding="utf-8")
    doc=Document();configure_styles(doc);configure_section(doc.sections[0],spec["running"]);add_cover(doc,spec);add_toc_field(doc);metrics=render_markdown(doc,markdown);add_side_story_legend(doc)
    assert_no_known_backstage_leak(visible_docx_text(doc))
    doc.core_properties.title=spec["title"];doc.core_properties.subject="Run26 from-scratch historical reader";doc.core_properties.author="Projet tourisme-etude-historico-geographique";doc.core_properties.comments="Rendered from reviewed structured-evidence manuscript; no previous reader DOCX used as input."
    output=root/spec["output"];doc.save(output)
    return {"project":key,"reviewed_paragraphs":reviewed,**metrics,"docx":str(output.relative_to(REPO)),"manuscript":str(manuscript.relative_to(REPO))}


def main()->int:
    p=argparse.ArgumentParser();p.add_argument("--project",choices=["pre","post","all"],default="all");a=p.parse_args();keys=list(SPECS) if a.project=="all" else [a.project];metrics=[build(k) for k in keys]
    path=REPO/"docs"/"RUN26_FROM_SCRATCH_RENDER_METRICS.json";path.write_text(json.dumps(metrics,ensure_ascii=False,indent=2)+"\n",encoding="utf-8");print(json.dumps(metrics,ensure_ascii=False,indent=2));return 0

if __name__=="__main__":raise SystemExit(main())
