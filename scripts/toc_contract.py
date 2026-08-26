#!/usr/bin/env python3
"""Structured table-of-contents helpers; no manuscript prose scan."""
from __future__ import annotations
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from build_story_scaffold import build_toc_from_scaffold


def inject_word_toc(doc, scaffold:dict, after_paragraph_index:int)->int:
    """Inject a real Word TOC field after the cover, using scaffold only as hierarchy gate."""
    toc=build_toc_from_scaffold(scaffold)
    if not toc:
        raise RuntimeError("story scaffold has no arcs; refusing to inject empty TOC")
    anchor=doc.paragraphs[after_paragraph_index]
    heading=doc.add_paragraph();run=heading.add_run("Sommaire");run.bold=True
    field_p=doc.add_paragraph()
    field=OxmlElement("w:fldSimple");field.set(qn("w:instr"),'TOC \\o "1-3" \\h \\z \\u');field_p._p.append(field)
    anchor._p.addnext(heading._p);heading._p.addnext(field_p._p)
    settings=doc.settings.element
    update=OxmlElement("w:updateFields");update.set(qn("w:val"),"true");settings.append(update)
    return len(toc)
