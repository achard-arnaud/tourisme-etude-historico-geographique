#!/usr/bin/env python3
"""Deterministic frontstage lint for rendered Word readers.

This catches machine-readable residue before visual QA.  It intentionally does
not claim to replace page-by-page inspection of pagination and typography.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document

from frontstage_reader_contract import assert_no_known_backstage_leak, visible_docx_text


RAW_TOKENS = (
    "<table>",
    "</table>",
    "<colgroup>",
    "<thead>",
    "<tbody>",
    "[ILLUSTRATION:",
    "<!-- [SIDE-STORY:",
    "Complément V3",
)
LINE_PATTERNS = (
    re.compile(r"^ARC A\d{2}\b", re.MULTILINE),
    re.compile(r"^HIL-[0-9/]+\s+[—–-]", re.MULTILINE),
    re.compile(r"^Z[0-9]+(?:/Z?[0-9]+)*\s+[—–-]", re.MULTILINE),
)


def lint(path: Path) -> dict[str, int | str]:
    doc = Document(path)
    visible = visible_docx_text(doc)
    errors = [f"visible raw token: {token}" for token in RAW_TOKENS if token.casefold() in visible.casefold()]
    errors.extend(f"visible production label: {pattern.pattern}" for pattern in LINE_PATTERNS if pattern.search(visible))
    try:
        assert_no_known_backstage_leak(visible)
    except RuntimeError as exc:
        errors.append(str(exc))
    if errors:
        raise RuntimeError(f"{path}: " + "; ".join(errors))
    return {
        "docx": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "visible_characters": len(visible),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", nargs="+", type=Path)
    args = parser.parse_args()
    for path in args.docx:
        print(lint(path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
