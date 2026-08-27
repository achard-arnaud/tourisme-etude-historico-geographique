#!/usr/bin/env python3
"""Final reader-facing cleanup after deterministic V3 composition.

The historical production process remains versioned in Git and metrics. This module
removes only known production apparatus that was previously inserted into the visible
reader (V1/V3/baseline/delta wording), then applies the side-story visual legend.

Run27 adds a two-stage return-target contract. The renderer itself never browses: any
missing canonical marker must have been resolved by a persisted research record first.
Supported research resolutions are materialised as hidden markers before layout; a
required return that still cannot resolve blocks final export.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from return_target_resolution import apply_project_research_resolutions, validate_required_return_targets
from side_story_presentation import LEGEND_HEADING, add_side_story_legend, apply_side_story_palette, markdown_legend


VISIBLE_REPLACEMENTS = {
    "ÉDITION V3 INTÉGRALE DE LECTURE — VOL RETOUR": "FRESQUE HISTORICO-GÉOGRAPHIQUE — VOLUME RETOUR",
    "V1 longue conservée, fiches de conversation et ajouts promus intégrés sans compression.": "Lecture intégrale : lieux, mécanismes, changements d’échelle et approfondissements de terrain.",
    "V3 intégrale — arcs chronologiques, HIL, comparateurs et zooms géographiques": "Lecture intégrale — arcs chronologiques, comparaisons et changements d’échelle",
    "État des données : 20 août 2026 · Usage personnel · Baseline V1 conservée": "Usage personnel · Sri Lanka · août 2026",
    "Appareil de sources des compléments V3": "Sources des approfondissements",
    "Ces notices donnent la source exacte des développements ajoutés à la V1. Le tier décrit la nature de la source, non une autorisation à généraliser au-delà de son champ.": "Ces notices donnent la source exacte des approfondissements. Le niveau de source décrit la nature du document, sans autoriser une généralisation au-delà de son champ.",
}

KNOWN_BACKSTAGE_FRAGMENTS = (
    "Politique éditoriale de la V3 intégrale",
    "V1 longue conservée",
    "Baseline V1 conservée",
    "Complément V3",
    "Appareil de sources des compléments V3",
    "Le petit report.md est traité comme un delta",
    "side-story lineage",
)


def _plain_markdown_line(line: str) -> str:
    line = re.sub(r"^#{1,6}\s+", "", line.strip())
    line = re.sub(r"[*_`]", "", line)
    return re.sub(r"\s+", " ", line).strip()


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def strip_method_block_from_docx(doc: Document, method_block: str) -> int:
    targets = {_plain_markdown_line(line) for line in method_block.splitlines() if _plain_markdown_line(line)}
    removed = 0
    for paragraph in list(doc.paragraphs):
        if _plain_markdown_line(paragraph.text) in targets:
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def clean_visible_docx_text(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        replacement = VISIBLE_REPLACEMENTS.get(text.strip())
        if replacement is not None:
            if paragraph.runs:
                paragraph.runs[0].text = replacement
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(replacement)
            changed += 1
            continue
        if "Complément V3 —" in text:
            for run in paragraph.runs:
                if "Complément V3 —" in run.text:
                    run.text = run.text.replace("Complément V3 —", "Approfondissement —")
                    changed += 1
                    break
    return changed


def strip_method_block_from_markdown(markdown: str, method_block: str) -> str:
    block = method_block.strip()
    if block:
        markdown = markdown.replace(block, "")
    markdown = markdown.replace("Complément V3 —", "Approfondissement —")
    for old, new in VISIBLE_REPLACEMENTS.items():
        markdown = markdown.replace(old, new)
    markdown = re.sub(r"\n{3,}", "\n\n", markdown)
    return markdown.strip() + "\n"


def visible_docx_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def assert_no_known_backstage_leak(text: str) -> None:
    normalized = text.replace("`", "")
    found = [fragment for fragment in KNOWN_BACKSTAGE_FRAGMENTS if fragment.casefold() in normalized.casefold()]
    if found:
        raise RuntimeError(f"reader-facing backstage leakage remains: {found}")


def finalize_reader(project: Path, spec: dict) -> dict:
    output_dir = project / "09_output"
    docx_path = output_dir / spec["output"]
    markdown_path = output_dir / spec["markdown_output"]

    # Stage 1: read the canonical Markdown before frontstage cleanup. Existing explicit
    # claim/bridge/arc markers resolve directly. Stage 2: persisted research resolutions
    # may add a hidden marker to a paragraph whose historical proposition was independently
    # supported. The renderer itself remains network-free and deterministic.
    canonical_markdown = markdown_path.read_text(encoding="utf-8")
    canonical_markdown, research = apply_project_research_resolutions(project, canonical_markdown)
    if research["errors"]:
        raise RuntimeError("return-target research resolution failed:\n- " + "\n- ".join(research["errors"]))
    return_errors, return_report = validate_required_return_targets(project, canonical_markdown)
    if return_errors:
        raise RuntimeError("required side-story return resolution failed:\n- " + "\n- ".join(return_errors))

    doc = Document(docx_path)
    removed = strip_method_block_from_docx(doc, spec.get("method_block", ""))
    replacements = clean_visible_docx_text(doc)
    palette = apply_side_story_palette(doc, project, canonical_markdown=canonical_markdown)
    legend_rows = add_side_story_legend(doc)
    assert_no_known_backstage_leak(visible_docx_text(doc))
    doc.save(docx_path)

    markdown = strip_method_block_from_markdown(canonical_markdown, spec.get("method_block", ""))
    if f"## {LEGEND_HEADING}" not in markdown:
        markdown = markdown.rstrip() + "\n\n" + markdown_legend() + "\n"
    assert_no_known_backstage_leak(markdown)
    markdown_path.write_text(markdown, encoding="utf-8")

    return {
        "backstage_paragraphs_removed": removed,
        "frontstage_replacements": replacements,
        "return_research_markers_materialized": len(research["applied"]),
        "return_research_challenged_or_redirected": research["challenged_or_redirected"],
        "required_returns_checked": len(return_report),
        "required_returns_unresolved": 0,
        "side_story_headers_styled": palette["headers_styled"],
        "side_story_body_paragraphs_styled": palette["body_paragraphs_styled"],
        "side_story_blocks_resolved": palette["resolved_blocks"],
        "side_story_return_markers": palette["return_markers"],
        "side_story_kinds_seen": palette["kinds_seen"],
        "side_story_legend_rows": legend_rows,
    }
