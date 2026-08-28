#!/usr/bin/env python3
"""Materialize the Run34 ancient-reader patch after frontstage cleanup.

The user-facing reader stays free of technical claim IDs. A separate instrumented
Markdown copy preserves real [claim:*] and [SIDE-STORY:*] markers for return-target,
coverage and reciprocal-lineage QA.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from side_story_contract import load_side_stories
from side_story_presentation import _decorate_paragraph, ensure_side_story_styles

CLAIM_MARKER = re.compile(r"\s*\[claim:[^\]]+\]", re.I)
SIDE_STORY_SENTINEL = "SS-R34-MUDRA-DETOUR-001"
INSTRUMENTED_NAME = "report_v3_full_run34_instrumented.md"


def load_patch(project: Path) -> dict[str, Any] | None:
    path = project / "09_output" / "run34_reader_patch.json"
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _stories(project: Path) -> dict[str, dict[str, Any]]:
    return {str(item["id"]): item for _, item in load_side_stories(project) if item.get("id")}


def _side_story_markdown(item: dict[str, Any], *, instrumented: bool) -> str:
    label = (item.get("render") or {}).get("label", "")
    title = item.get("title", "")
    body = (item.get("content") or {}).get("body_markdown", "").strip()
    visible = f"**{label} — {title}**\n\n{body}"
    if not instrumented:
        return visible
    sid = item["id"]
    return (
        f"<!-- [SIDE-STORY:{sid}] BEGIN kind={item.get('kind','')} -->\n"
        f"{visible}\n"
        f"<!-- [SIDE-STORY:{sid}] END -->"
    )


def build_markdown_block(project: Path, patch: dict[str, Any], *, instrumented: bool) -> str:
    stories = _stories(project)
    chunks: list[str] = []
    for row in patch.get("sequence") or []:
        if row.get("type") == "paragraph":
            text = str(row.get("text") or "").strip()
            chunks.append(text if instrumented else CLAIM_MARKER.sub("", text).strip())
        elif row.get("type") == "side_story":
            sid = str(row.get("id") or "")
            if sid not in stories:
                raise RuntimeError(f"Run34 reader patch references unknown side story: {sid}")
            chunks.append(_side_story_markdown(stories[sid], instrumented=instrumented))
        else:
            raise RuntimeError(f"Run34 reader patch has unsupported row: {row}")
    return "\n\n".join(chunks)


def _normalize_heading(line: str) -> str:
    return re.sub(r"[*_`]", "", re.sub(r"^#+\s*", "", line)).strip()


def insert_after_markdown_anchor(text: str, anchor: str, block: str) -> str:
    if SIDE_STORY_SENTINEL in text:
        return text
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if _normalize_heading(line).startswith(anchor):
            lines[index + 1:index + 1] = ["", block.strip(), ""]
            return "\n".join(lines).rstrip() + "\n"
    raise RuntimeError(f"Run34 reader anchor missing from markdown: {anchor}")


def _find_docx_anchor(doc: Document, anchor: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(anchor):
            return paragraph
    raise RuntimeError(f"Run34 reader anchor missing from docx: {anchor}")


def _append_plain_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.add_run(CLAIM_MARKER.sub("", text).strip())
    return paragraph


def _append_side_story(doc: Document, item: dict[str, Any]):
    kind = str(item.get("kind"))
    label = (item.get("render") or {}).get("label", "")
    title = item.get("title", "")
    body = (item.get("content") or {}).get("body_markdown", "").strip()
    header = doc.add_paragraph()
    run = header.add_run(f"{label} — {title}")
    run.bold = True
    body_p = doc.add_paragraph(body)
    _decorate_paragraph(header, kind, header=True, position="start")
    _decorate_paragraph(body_p, kind, header=False, position="end")
    return [header, body_p]


def patch_docx(project: Path, docx_path: Path, patch: dict[str, Any]) -> int:
    doc = Document(docx_path)
    if any(SIDE_STORY_SENTINEL in p.text for p in doc.paragraphs):
        return 0
    ensure_side_story_styles(doc)
    anchor = _find_docx_anchor(doc, str(patch["anchor"]))
    stories = _stories(project)
    paragraphs = []
    side_story_count = 0
    for row in patch.get("sequence") or []:
        if row.get("type") == "paragraph":
            paragraphs.append(_append_plain_paragraph(doc, str(row.get("text") or "")))
        elif row.get("type") == "side_story":
            sid = str(row.get("id") or "")
            if sid not in stories:
                raise RuntimeError(f"Run34 reader patch references unknown side story: {sid}")
            paragraphs.extend(_append_side_story(doc, stories[sid]))
            side_story_count += 1
        else:
            raise RuntimeError(f"Run34 reader patch has unsupported row: {row}")
    for paragraph in reversed(paragraphs):
        anchor._p.addnext(paragraph._p)
    doc.save(docx_path)
    return side_story_count


def apply_run34_reader_patch(project: Path, spec: dict[str, Any]) -> dict[str, Any]:
    patch = load_patch(project)
    if patch is None:
        return {"run34_reader_patch": "absent"}
    output_dir = project / "09_output"
    reader_md = output_dir / spec["markdown_output"]
    reader_docx = output_dir / spec["output"]
    original = reader_md.read_text(encoding="utf-8")
    reader_block = build_markdown_block(project, patch, instrumented=False)
    instrumented_block = build_markdown_block(project, patch, instrumented=True)
    reader_text = insert_after_markdown_anchor(original, str(patch["anchor"]), reader_block)
    instrumented_text = insert_after_markdown_anchor(original, str(patch["anchor"]), instrumented_block)
    reader_md.write_text(reader_text, encoding="utf-8")
    instrumented_path = output_dir / INSTRUMENTED_NAME
    instrumented_path.write_text(instrumented_text, encoding="utf-8")
    boxes = patch_docx(project, reader_docx, patch)
    return {
        "run34_reader_patch": "materialized",
        "run34_side_stories_patched": boxes,
        "run34_instrumented_markdown": str(instrumented_path),
    }


if __name__ == "__main__":
    raise SystemExit("run34_reader_patch.py is invoked by render_composed_reader.py")
