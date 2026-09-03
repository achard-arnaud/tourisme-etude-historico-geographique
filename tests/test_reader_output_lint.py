import tempfile
import unittest
from pathlib import Path
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from lint_reader_output import lint


class ReaderOutputLintTests(unittest.TestCase):
    def test_accepts_clean_reader(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "clean.docx"
            doc = Document()
            doc.add_heading("Chapitre 1 — Ouverture", 1)
            doc.add_paragraph("Un paragraphe destiné au lecteur.")
            doc.save(path)
            self.assertEqual(2, lint(path)["paragraphs"])

    def test_rejects_visible_raw_html(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "bad.docx"
            doc = Document()
            doc.add_paragraph("<table>")
            doc.save(path)
            with self.assertRaises(RuntimeError):
                lint(path)


if __name__ == "__main__":
    unittest.main()
