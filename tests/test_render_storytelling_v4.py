import tempfile
import unittest
from pathlib import Path
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from scripts.render_storytelling_v4 import build


class RenderStorytellingV4Tests(unittest.TestCase):
    def test_renders_legacy_headings_tables_and_side_story(self):
        markdown = """# **Sommaire**

**PARTIE I — TEST**

> Chapitre 1 — Test

**PARTIE I — TEST**

**Chapitre 1 — Test**

## Section

| **A** | **B** |
|---|---|
| alpha | beta |

<!-- [SIDE-STORY:SS-TEST] BEGIN kind=detour -->
**Détour — Test**

Corps du détour.
<!-- [SIDE-STORY:SS-TEST] END -->

<table>
<thead><tr><th><p><strong>PETIT DÉTOUR</strong></p><p>Encadré HTML.</p></th></tr></thead>
</table>
"""
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "input.md"
            output = root / "output.docx"
            source.write_text(markdown, encoding="utf-8")
            metrics = build(source, output)
            doc = Document(output)
            self.assertTrue(output.exists())
            self.assertEqual(1, metrics["tables"])
            self.assertEqual(1, metrics["side_story_blocks"])
            self.assertEqual(1, metrics["html_callouts"])
            self.assertIn("Chapitre 1 — Test", [p.text for p in doc.paragraphs])
            self.assertGreaterEqual(len(doc.tables), 4)  # cover, Markdown table, two side stories


if __name__ == "__main__":
    unittest.main()
