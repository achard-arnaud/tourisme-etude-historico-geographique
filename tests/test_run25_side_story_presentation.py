import sys,unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
sys.path.insert(0,str(ROOT/"scripts"))

from frontstage_reader_contract import (
    assert_no_known_backstage_leak,
    clean_visible_docx_text,
    strip_method_block_from_docx,
    strip_method_block_from_markdown,
)
from paragraph_review_gate import review_paragraph
from side_story_contract import KINDS,RENDER_LABELS
from side_story_presentation import (
    KIND_ORDER,
    LEGEND_HEADING,
    SIDE_STORY_PRESENTATION,
    add_side_story_legend,
    apply_side_story_palette,
    detect_kind,
)


class Run25SideStoryPaletteTests(unittest.TestCase):
    def test_palette_covers_every_contract_kind_exactly(self):
        self.assertEqual(KINDS,set(SIDE_STORY_PRESENTATION))
        self.assertEqual(10,len(KIND_ORDER))
        self.assertEqual(10,len({SIDE_STORY_PRESENTATION[k]["symbol"] for k in KIND_ORDER}))
        expected={
            "false_lead":"FFF6D8","detour":"E3EEF7","dezoom":"E7EBF0","also":"E7F3E8","method":"F0EDE7",
            "portrait":"F7E7E1","object_focus":"F5EFDD","comparator":"EEE7F5","callback":"FBE6E3","analytical_focus":"DCE6EE",
        }
        self.assertEqual(expected,{k:SIDE_STORY_PRESENTATION[k]["fill"] for k in KIND_ORDER})

    def test_kind_detection_survives_symbol_and_case(self):
        self.assertEqual("false_lead",detect_kind("① FAUSSE PISTE — Et si Kandy était sans défense ?"))
        self.assertEqual("method",detect_kind("Point de méthode — dater une inscription"))
        self.assertIsNone(detect_kind("Politique éditoriale de la V3 intégrale"))

    def test_all_headers_receive_redundant_symbol_and_fill(self):
        doc=Document()
        for kind in KIND_ORDER:
            doc.add_paragraph(f"{RENDER_LABELS[kind]} — Exemple")
        stats=apply_side_story_palette(doc)
        self.assertEqual(10,stats["headers_styled"])
        for paragraph,kind in zip(doc.paragraphs,KIND_ORDER):
            self.assertTrue(paragraph.text.startswith(SIDE_STORY_PRESENTATION[kind]["symbol"]))
            shd=paragraph._p.get_or_add_pPr().find(qn("w:shd"))
            self.assertIsNotNone(shd)
            self.assertEqual(SIDE_STORY_PRESENTATION[kind]["fill"],shd.get(qn("w:fill")))

    def test_legend_maps_all_ten_symbols_and_method_distinction(self):
        doc=Document();rows=add_side_story_legend(doc)
        self.assertEqual(10,rows)
        self.assertTrue(any(p.text==LEGEND_HEADING for p in doc.paragraphs))
        text="\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        for kind in KIND_ORDER:
            self.assertIn(SIDE_STORY_PRESENTATION[kind]["symbol"],text)
            self.assertIn(RENDER_LABELS[kind],text)
        visible="\n".join(p.text for p in doc.paragraphs)
        self.assertIn("méthode historique",visible)
        self.assertIn("ne désigne jamais le processus de production",visible)


class Run25FrontstageCleanupTests(unittest.TestCase):
    METHOD_BLOCK="""## Politique éditoriale de la V3 intégrale

Cette édition prend la V1 comme baseline non destructible.
Le petit `report.md` est traité comme un delta promu.
"""

    def test_reader_method_side_story_is_allowed_but_production_method_is_not(self):
        claim={"id":"C-LEGACY","claim":"x"}
        good=review_paragraph("Point de méthode — Une inscription se date d'abord par sa paléographie et son contexte.",claim=claim)
        self.assertTrue(good.passed,good.violations)
        bad=review_paragraph("Politique éditoriale de la V3 intégrale : la baseline conserve le delta.",claim=claim)
        self.assertFalse(bad.passed)
        self.assertEqual("Don't",bad.violations[0].category)

    def test_method_block_is_removed_from_markdown_not_reworded_as_history(self):
        markdown="# Livre\n\n## Méthode de lecture\n\n"+self.METHOD_BLOCK+"\n## 1. Histoire\n\nTexte.\n"
        cleaned=strip_method_block_from_markdown(markdown,self.METHOD_BLOCK)
        self.assertNotIn("Politique éditoriale",cleaned)
        self.assertNotIn("report.md",cleaned)
        self.assertIn("## 1. Histoire",cleaned)

    def test_docx_cleanup_removes_method_block_and_rewrites_version_labels(self):
        doc=Document()
        doc.add_paragraph("ÉDITION V3 INTÉGRALE DE LECTURE — VOL RETOUR")
        doc.add_heading("Politique éditoriale de la V3 intégrale",level=2)
        doc.add_paragraph("Cette édition prend la V1 comme baseline non destructible.")
        doc.add_paragraph("Le petit `report.md` est traité comme un delta promu.")
        doc.add_heading("Complément V3 — Kandy",level=2)
        removed=strip_method_block_from_docx(doc,self.METHOD_BLOCK)
        changed=clean_visible_docx_text(doc)
        text="\n".join(p.text for p in doc.paragraphs)
        self.assertEqual(3,removed)
        self.assertGreaterEqual(changed,2)
        self.assertIn("FRESQUE HISTORICO-GÉOGRAPHIQUE — VOLUME RETOUR",text)
        self.assertIn("Approfondissement — Kandy",text)
        assert_no_known_backstage_leak(text)


if __name__=="__main__":unittest.main()
