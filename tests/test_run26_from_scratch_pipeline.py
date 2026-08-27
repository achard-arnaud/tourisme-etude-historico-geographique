import json,sys,tempfile,unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
sys.path.insert(0,str(ROOT/"scripts"))

from build_from_scratch_packets import ReadLedger,build_packets
from from_scratch_review_contract import initialize_ledger,validate_review_ledger
from materialize_side_stories import materialize_text,side_story_begin_marker,side_story_end_marker
from render_from_scratch_reader import add_side_story_block
from sarah_voice_contract import review_skeleton
from side_story_presentation import SIDE_STORY_PRESENTATION


def passing_style_review(text):
    record=review_skeleton(text,generation_pass_id="gen-pass-1",generation_context_id="ctx-gen-1")
    record.update({"passed":True,"evaluator":"bounded_llm","review_pass_id":"style-pass-1","review_context_id":"ctx-style-1"})
    record["marker_results"]={
        "scope_precision":{"status":"pass","rationale":"La portée reste bornée au cas décrit."},
        "continuous_prose_not_social_format":{"status":"pass","rationale":"La prose n'imite pas un format social."},
        "lived_opening_callback":{"status":"not_applicable","rationale":"Cette unité n'est pas une ouverture disposant de matière vécue."},
        "rigor_compressed_in_sentence":{"status":"not_applicable","rationale":"Cette unité ne porte pas de réserve méthodologique particulière."},
        "concrete_texture_before_abstraction":{"status":"pass","rationale":"Le fait concret précède la conclusion."},
    }
    return record


class Run26PacketTests(unittest.TestCase):
    def test_read_ledger_blocks_reader_prose_and_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp);(project/"09_output/archive").mkdir(parents=True)
            report=project/"09_output/report_v3_full.md";report.write_text("old prose",encoding="utf-8")
            docx=project/"old.docx";docx.write_bytes(b"x")
            ledger=ReadLedger(project)
            with self.assertRaisesRegex(RuntimeError,"contamination blocked"):ledger.text(report)
            with self.assertRaisesRegex(RuntimeError,"contamination blocked"):ledger.text(docx)

    def test_packet_builder_uses_structured_evidence_and_initial_false_state(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp)/"p";(project/"01_arcs/A1/claims").mkdir(parents=True);(project/"05_sources").mkdir();(project/"06_bridges").mkdir();(project/"08_questions").mkdir();(project/"09_output/side_stories").mkdir(parents=True);(project/"00_method/capture").mkdir(parents=True)
            (project/"01_arcs/A1/ARC.md").write_text("# A1\n\nArc evidence note.",encoding="utf-8")
            claim={"id":"C1","claim":"Fact","hil":"HIL-02_geography-environment","source_ids":["S1"]}
            (project/"01_arcs/A1/claims/C1.json").write_text(json.dumps(claim),encoding="utf-8")
            (project/"05_sources/source_register.json").write_text(json.dumps([{"id":"S1","title":"Source"}]),encoding="utf-8")
            output=project/"packet_out";manifest=build_packets(project,output);packet=json.loads((output/"A1.json").read_text(encoding="utf-8"))
            self.assertFalse(manifest["contamination_check"]["reader_prose_loaded"])
            self.assertEqual(["C1"],[x["id"] for x in packet["claims"]])
            self.assertIn("HIL-02_geography-environment",packet["relevant_hil"])
            self.assertEqual({"checklist_reviewed":False,"sarah_style_reviewed":False,"hil_scope_reviewed":False},packet["drafting_contract"]["paragraph_review_state_initial"])
            self.assertTrue(all("09_output/report" not in p for p in manifest["read_ledger"]))


class Run26ReviewLedgerTests(unittest.TestCase):
    def test_ledger_initializes_all_review_flags_false_and_binds_voice_contract(self):
        records=initialize_ledger("# T\n\nUn paragraphe. [claim:C1]\n",generation_pass_id="gen-pass-1",generation_context_id="ctx-gen-1")
        self.assertEqual(1,len(records));self.assertEqual({"checklist_reviewed":False,"sarah_style_reviewed":False,"hil_scope_reviewed":False},records[0]["initial_state"]);self.assertEqual(records[0]["initial_state"],records[0]["review_state"])
        style=records[0]["sarah_style_review"];self.assertFalse(style["passed"]);self.assertTrue(style["paragraph_sha256"]);self.assertEqual("sarah-voice-run25-v1",style["voice_contract_id"])

    def test_final_ledger_rejects_irrelevant_hil(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp);(project/"01_arcs/A1/claims").mkdir(parents=True)
            (project/"01_arcs/A1/claims/C1.json").write_text(json.dumps({"id":"C1","claim":"x","hil":"HIL-06_security-coercion"}),encoding="utf-8")
            markdown="# T\n\nUn paragraphe. [claim:C1]\n";record=initialize_ledger(markdown)[0]
            paragraph="Un paragraphe. [claim:C1]"
            record["review_state"]={"checklist_reviewed":True,"sarah_style_reviewed":True,"hil_scope_reviewed":True};record["selected_hil_ids"]=["HIL-03_economy-infrastructure"];record["sarah_style_review"]=passing_style_review(paragraph)
            errors=validate_review_ledger(project,markdown,[record]);self.assertTrue(any("irrelevant HIL" in e for e in errors))

    def test_complete_review_ledger_accepts_relevant_subset(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp);(project/"01_arcs/A1/claims").mkdir(parents=True)
            (project/"01_arcs/A1/claims/C1.json").write_text(json.dumps({"id":"C1","claim":"x","hil_ids":["HIL-06_security-coercion","HIL-02_geography-environment"]}),encoding="utf-8")
            markdown="# T\n\nUn paragraphe. [claim:C1]\n";record=initialize_ledger(markdown)[0]
            paragraph="Un paragraphe. [claim:C1]"
            record["review_state"]={"checklist_reviewed":True,"sarah_style_reviewed":True,"hil_scope_reviewed":True};record["selected_hil_ids"]=["HIL-06_security-coercion"];record["sarah_style_review"]=passing_style_review(paragraph)
            self.assertEqual([],validate_review_ledger(project,markdown,[record]))

    def test_ledger_rejects_stale_sarah_review_after_text_change(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp);(project/"01_arcs/A1/claims").mkdir(parents=True)
            (project/"01_arcs/A1/claims/C1.json").write_text(json.dumps({"id":"C1","claim":"x"}),encoding="utf-8")
            markdown="# T\n\nTexte réécrit. [claim:C1]\n";record=initialize_ledger(markdown)[0]
            record["review_state"]={"checklist_reviewed":True,"sarah_style_reviewed":True,"hil_scope_reviewed":True};record["sarah_style_review"]=passing_style_review("Ancien texte. [claim:C1]")
            errors=validate_review_ledger(project,markdown,[record]);self.assertTrue(any("paragraph hash is stale" in e for e in errors))


class Run26SideStoryBoundaryTests(unittest.TestCase):
    def test_materializer_emits_balanced_begin_end_fences(self):
        with tempfile.TemporaryDirectory() as tmp:
            project=Path(tmp);(project/"09_output/side_stories").mkdir(parents=True)
            item={"schema_version":"1.1","class":"side_story","id":"SS-X","kind":"detour","status":"promoted","title":"Example","purpose":"p","reason_off_trunk":"r","payoff":"p","map_eligible":False,"reader_policy":{"min_age":10},"lineage":{"claim_ids":[],"source_ids":[],"bridge_ids":[],"hil_ids":[],"drift_paths":[],"origin_paths":[]},"placement":{"section_anchor":"Anchor","return_to":"anchor:Body"},"zoom_excursion":None,"content":{"takeaway":"A distinct useful conclusion.","body_markdown":"Inserted body"},"render":{"label":"Petit détour","marker":"[SIDE-STORY:SS-X]","required_in_reader":False}}
            (project/"09_output/side_stories/SS-X.json").write_text(json.dumps(item),encoding="utf-8")
            text,count=materialize_text(project,"# A\n\nAnchor\n\nBody\n");self.assertEqual(1,count);self.assertIn(side_story_begin_marker(item),text);self.assertIn(side_story_end_marker(item),text);self.assertLess(text.index("BEGIN"),text.index("Inserted body"));self.assertLess(text.index("Inserted body"),text.index(" END"))

    def test_from_scratch_side_story_is_one_filled_closed_cell(self):
        doc=Document();add_side_story_block(doc,"false_lead",["**Fausse piste — Exemple**","La réponse reste dans le même encadré."])
        self.assertEqual(1,len(doc.tables));cell=doc.tables[0].cell(0,0);tc_pr=cell._tc.get_or_add_tcPr();shd=tc_pr.find(qn("w:shd"));self.assertEqual(SIDE_STORY_PRESENTATION["false_lead"]["fill"],shd.get(qn("w:fill")))
        borders=tc_pr.find(qn("w:tcBorders"));self.assertIsNotNone(borders)
        for edge in ("top","bottom","left","right"):self.assertIsNotNone(borders.find(qn(f"w:{edge}")))
        self.assertIn("①",cell.text);self.assertIn("La réponse",cell.text)


if __name__=="__main__":unittest.main()
